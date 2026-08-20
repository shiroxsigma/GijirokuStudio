import os
import re
import sys
import json
import time
import queue
import datetime
import threading
import traceback
import subprocess
import base64
import html
import shutil
import tempfile
import urllib.request
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog

import numpy as np
import mss
from PIL import Image, ImageTk
import imagehash
import sounddevice as sd
import pyaudiowpatch as pyaudio

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FFMPEG_PATH = os.path.join(BASE_DIR, "ffmpeg.exe")
SETTINGS_PATH = os.path.join(BASE_DIR, "settings.json")
GLOSSARY_PATH = os.path.join(BASE_DIR, "glossary.csv")

# Extra entry in the display picker: capture a dragged rectangle instead of a
# whole monitor.
REGION_CHOICE = "範囲を指定（ドラッグ）"
MIN_REGION = 40          # physical px; smaller selections are almost certainly slips

SUMMARY_PROVIDERS = [
    ("なし（要約しない）", "none"),
    ("Ollama（ローカル完結）", "ollama"),
    ("Claude API（外部送信）", "claude"),
]

LANG_OPTIONS = [
    ("自動（日本語/英語）", "auto"),
    ("日本語", "ja"), ("English", "en"), ("中文", "zh"), ("한국어", "ko"),
]

# Sentinel placed on a capture queue to signal end-of-stream to the mixer.
_EOF = object()

# Per-speaker tracks. Mics are "me", speaker loopbacks are "them" — the two
# roles are already physically separate, so no diarization model is needed.
ROLE_SELF = "自分"
ROLE_OTHER = "相手"
ROLE_TRACK_SELF = "audio_self.mp3"
ROLE_TRACK_OTHER = "audio_other.mp3"


class _GlobalHotkey:
    """A system-wide hotkey, registered with Win32 on its own thread.

    tkinter's `bind` only fires while the app has focus, but during a meeting
    the foreground window is Teams or Zoom — so marking an important moment
    needs a hotkey the OS delivers no matter what is on top. RegisterHotKey
    binds to the calling thread's message queue, so registration and the
    message pump have to live on the same thread.
    """

    MODIFIERS = {"alt": 0x0001, "ctrl": 0x0002, "control": 0x0002,
                 "shift": 0x0004, "win": 0x0008}
    MOD_NOREPEAT = 0x4000
    WM_HOTKEY = 0x0312
    WM_QUIT = 0x0012

    def __init__(self, spec, callback, log=print):
        self.spec = spec
        self.callback = callback
        self.log = log
        self.mods, self.vk = self._parse(spec)
        self._thread = None
        self._thread_id = None
        self._ready = threading.Event()
        self.registered = False

    @classmethod
    def _parse(cls, spec):
        """'ctrl+shift+m' -> (modifier mask, virtual key code)."""
        mods = cls.MOD_NOREPEAT
        vk = None
        for part in str(spec).lower().split("+"):
            part = part.strip()
            if not part:
                continue
            if part in cls.MODIFIERS:
                mods |= cls.MODIFIERS[part]
            elif len(part) == 1:
                vk = ord(part.upper())
            elif part.startswith("f") and part[1:].isdigit():
                vk = 0x70 + int(part[1:]) - 1      # VK_F1 == 0x70
        return mods, vk

    def start(self):
        if os.name != "nt" or self.vk is None:
            self.log(f"[ホットキー] 使用できません: {self.spec}")
            return False
        self._thread = threading.Thread(target=self._run, daemon=True)
        self._thread.start()
        self._ready.wait(timeout=2)
        return self.registered

    def _run(self):
        import ctypes
        from ctypes import wintypes
        user32 = ctypes.windll.user32
        self._thread_id = ctypes.windll.kernel32.GetCurrentThreadId()
        hotkey_id = 1
        try:
            ok = user32.RegisterHotKey(None, hotkey_id, self.mods, self.vk)
        except Exception as e:
            self.log(f"[ホットキー] 登録失敗: {e}")
            self._ready.set()
            return
        if not ok:
            # Almost always means another application already owns the combo.
            self.log(f"[ホットキー] {self.spec} は他のアプリが使用中のため登録できません")
            self._ready.set()
            return
        self.registered = True
        self._ready.set()
        try:
            msg = wintypes.MSG()
            while user32.GetMessageW(ctypes.byref(msg), None, 0, 0) > 0:
                if msg.message == self.WM_HOTKEY and msg.wParam == hotkey_id:
                    try:
                        self.callback()
                    except Exception as e:
                        print(f"[ホットキー処理警告] {e}")
        finally:
            user32.UnregisterHotKey(None, hotkey_id)
            self.registered = False

    def stop(self):
        if self._thread is None or self._thread_id is None:
            return
        try:
            import ctypes
            ctypes.windll.user32.PostThreadMessageW(
                self._thread_id, self.WM_QUIT, 0, 0)
        except Exception as e:
            print(f"[ホットキー終了警告] {e}")
        self._thread = None


class _PcmWriter:
    """An ffmpeg encoder fed stereo s16le PCM on stdin (the Captura pattern).

    Role tracks pass `transcription_only`: they are downmixed to 16kHz mono at
    64kbps because nothing ever listens to them — they exist to be handed to
    whisper — while audio_main.mp3 stays full quality. Both are fed the exact
    same PCM, so ffmpeg's resampling keeps them on one timeline.
    """

    def __init__(self, path, rate, transcription_only=False):
        self.path = path
        self.written = 0
        args = [FFMPEG_PATH, "-f", "s16le", "-acodec", "pcm_s16le",
                "-ar", str(rate), "-ac", "2", "-i", "-", "-c:a", "libmp3lame"]
        if transcription_only:
            args += ["-ar", "16000", "-ac", "1", "-b:a", "64k"]
        else:
            args += ["-b:a", "192k"]
        args += ["-y", path]
        flags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
        self.proc = subprocess.Popen(
            args, stdin=subprocess.PIPE, stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL, creationflags=flags)

    def write(self, pcm_bytes):
        if self.proc is None or self.proc.poll() is not None:
            return
        try:
            self.proc.stdin.write(pcm_bytes)
            self.written += len(pcm_bytes)
        except (BrokenPipeError, OSError, ValueError):
            pass   # one dead encoder must not stop the others

    def close(self):
        if self.proc is None:
            return
        try:
            self.proc.stdin.close()
        except Exception as e:
            print(f"[ffmpeg終了警告] stdin ({os.path.basename(self.path)}): {e}")
        try:
            self.proc.wait(timeout=10)
        except subprocess.TimeoutExpired:
            self.proc.kill()
            try:
                self.proc.wait(timeout=5)
            except Exception as e:
                print(f"[ffmpeg終了警告] kill: {e}")
        except Exception as e:
            print(f"[ffmpeg終了警告] wait: {e}")
        self.proc = None


def _com_initialize():
    """Initialise COM on the calling thread; True if we own the initialisation.

    WASAPI is COM based and COM apartments are per thread. PortAudio only sets
    COM up on the thread that first calls Pa_Initialize — later calls are
    reference-counted no-ops — so a second capture thread inherits no apartment
    and Pa_StartStream fails there with "Unanticipated host error" (-9999) even
    though the device opened fine. Every capture thread initialises COM itself.
    """
    if os.name != "nt":
        return False
    try:
        import ctypes
        hr = ctypes.windll.ole32.CoInitializeEx(None, 0x2)  # APARTMENTTHREADED
    except Exception as e:
        print(f"[COM初期化警告] {e}")
        return False
    # S_OK / S_FALSE: this thread owns an initialisation and must balance it.
    # RPC_E_CHANGED_MODE and other failures: leave the existing apartment alone.
    return hr in (0, 1)


def _com_uninitialize():
    try:
        import ctypes
        ctypes.windll.ole32.CoUninitialize()
    except Exception as e:
        print(f"[COM終了警告] {e}")


class MeetingRecorderGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("GijirokuStudio v2 - リモート会議記録システム")
        self.root.geometry("620x860")
        self.root.resizable(False, False)

        self.is_recording = False
        self.stop_event = threading.Event()
        self.ffmpeg_proc = None       # _PcmWriter for the mixed audio
        self._role_writers = {}       # role -> _PcmWriter (speaker separation)
        self.transcriber = None
        self.transcription_file = None
        self._transcribe_start = 0
        self._record_start = 0
        self._audio_level = 0.0
        self._preloaded_transcriber = None
        self._recording_dir = None
        self._recording_mon_idx = 0
        self._recording_rect = None   # frozen at record start, like the dHash threshold
        self._pipeline_ctx = None
        self._pipeline_thread = None
        self._last_record_dir = None
        self._queue_overflow_count = 0
        self._meeting_name = ""
        self._rt_lang = None  # None = auto (ja/en detect), else fixed lang code
        self._post_cancel = threading.Event()
        self._post_running = False
        self._marker_lock = threading.Lock()
        self._marker_count = 0
        self._hotkey = None
        self._browser_win = None

        self._audio_devices = []        # list of enumerated device dicts
        self._audio_queues = []         # per-device capture queues (during recording)
        self.transcribe_queue = None
        # PortAudio init/open/terminate are not thread-safe: with several devices
        # starting at once, unrelated opens fail with bogus "invalid sample rate".
        self._open_lock = threading.Lock()
        self._pcm_written = 0           # bytes handed to ffmpeg this recording

        self.INTERVAL = 5.0
        self.DHASH_THRESHOLD = 10
        self.JPEG_QUALITY = 85
        self.AUDIO_GAIN = 2.0
        self.TARGET_RATE = 44100
        self.TRANSCRIBE_RATE = 16000
        self.TRANSCRIBE_CHUNK_SECONDS = 5.0   # real-time chunk length fed to faster-whisper
        self.TRANSCRIBE_OVERLAP_SECONDS = 1.0  # trailing overlap kept across chunks
        self.STARVE_SECONDS = 0.5   # a source silent this long stops blocking the mix

        # faster-whisper (post-processing) settings — editable via settings.json
        self.WHISPER_MODEL = "large-v3-turbo"
        self.WHISPER_DEVICE = "cpu"
        self.WHISPER_COMPUTE = "int8"

        # faster-whisper (real-time) settings — smaller model for low latency
        self.REALTIME_WHISPER_MODEL = "base"

        # System-wide key for marking an important moment mid-meeting
        self.MARKER_HOTKEY = "ctrl+shift+m"

        # Screen capture area: None = whole display, else an mss rect
        self.CAPTURE_REGION = None
        self.CAPTURE_MAX_EDGE = 1600   # 0 disables the downscale

        # Slide OCR during post-processing. Off by default: it needs the
        # Windows Japanese OCR language pack and adds time to every run.
        self.OCR_ENABLED = False

        # AI summary of the transcript, generated during post-processing
        self.SUMMARY_PROVIDER = "none"   # none | ollama | claude
        self.SUMMARY_MODEL = ""          # empty = provider default

        self._enumerate_audio_devices()
        self._load_settings()
        self.create_widgets()
        self._populate_audio_listbox()
        self._update_region_label()
        self._tick_level_meter()

    # --------------------------------------------------------- Settings persistence

    def _load_settings(self):
        try:
            if os.path.exists(SETTINGS_PATH):
                with open(SETTINGS_PATH, "r", encoding="utf-8") as f:
                    s = json.load(f)
                self.DHASH_THRESHOLD = int(s.get("dhash_threshold", self.DHASH_THRESHOLD))
                self.AUDIO_GAIN = float(s.get("audio_gain", self.AUDIO_GAIN))
                self.JPEG_QUALITY = int(s.get("jpeg_quality", self.JPEG_QUALITY))
                self.WHISPER_MODEL = str(s.get("whisper_model", self.WHISPER_MODEL))
                self.WHISPER_DEVICE = str(s.get("whisper_device", self.WHISPER_DEVICE))
                self.WHISPER_COMPUTE = str(s.get("whisper_compute", self.WHISPER_COMPUTE))
                self.REALTIME_WHISPER_MODEL = str(s.get(
                    "realtime_whisper_model", self.REALTIME_WHISPER_MODEL))
                self.MARKER_HOTKEY = str(s.get("marker_hotkey", self.MARKER_HOTKEY))
                self.CAPTURE_MAX_EDGE = int(s.get("capture_max_edge", self.CAPTURE_MAX_EDGE))
                region = s.get("capture_region")
                if isinstance(region, dict) and all(
                        k in region for k in ("left", "top", "width", "height")):
                    self.CAPTURE_REGION = {k: int(region[k])
                                           for k in ("left", "top", "width", "height")}
                self.OCR_ENABLED = bool(s.get("ocr_enabled", self.OCR_ENABLED))
                self.SUMMARY_PROVIDER = str(s.get("summary_provider", self.SUMMARY_PROVIDER))
                self.SUMMARY_MODEL = str(s.get("summary_model", self.SUMMARY_MODEL))
        except Exception as e:
            print(f"[設定読み込み警告] {e}")

    def _save_settings(self):
        try:
            # read-modify-write so manually-added keys (incl. whisper_*) are preserved
            existing = {}
            if os.path.exists(SETTINGS_PATH):
                try:
                    with open(SETTINGS_PATH, "r", encoding="utf-8") as f:
                        existing = json.load(f)
                except Exception:
                    existing = {}
            existing.update({
                "dhash_threshold": self.DHASH_THRESHOLD,
                "audio_gain": self.AUDIO_GAIN,
                "jpeg_quality": self.JPEG_QUALITY,
                "whisper_model": self.WHISPER_MODEL,
                "whisper_device": self.WHISPER_DEVICE,
                "whisper_compute": self.WHISPER_COMPUTE,
                "realtime_whisper_model": self.REALTIME_WHISPER_MODEL,
                "marker_hotkey": self.MARKER_HOTKEY,
                "capture_region": self.CAPTURE_REGION,
                "capture_max_edge": self.CAPTURE_MAX_EDGE,
                "ocr_enabled": self.OCR_ENABLED,
                "summary_provider": self.SUMMARY_PROVIDER,
                "summary_model": self.SUMMARY_MODEL,
            })
            with open(SETTINGS_PATH, "w", encoding="utf-8") as f:
                json.dump(existing, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"[設定保存警告] {e}")

    # ------------------------------------------------------------------ UI

    def create_widgets(self):
        # ---- Menu bar ----
        menubar = tk.Menu(self.root)
        menu_settings = tk.Menu(menubar, tearoff=0)
        menu_settings.add_command(label="設定を開く...", command=self._open_settings_dialog)
        menu_settings.add_separator()
        menu_settings.add_command(label="設定をリセット", command=self._reset_settings)
        menubar.add_cascade(label="設定", menu=menu_settings)
        menu_rec = tk.Menu(menubar, tearoff=0)
        menu_rec.add_command(label="記録一覧を開く...",
                             command=self._open_recordings_browser)
        menubar.add_cascade(label="記録", menu=menu_rec)
        self.root.config(menu=menubar)

        frame_info = ttk.LabelFrame(self.root, text=" システム概要 ", padding=10)
        frame_info.pack(fill="x", padx=15, pady=(8, 4))
        ttk.Label(frame_info, justify="left", text=(
            "・スピーカー/マイク音声をリアルタイムMP3録音（ffmpegエンコード）\n"
            "・5秒ごとに画面変化を検知し、スライド切替時のみJPEG保存\n"
            "・高負荷モード: ローカルAIでリアルタイム文字起こし"
        )).pack(anchor="w")

        frame_set = ttk.LabelFrame(self.root, text=" 録画設定 ", padding=10)
        frame_set.pack(fill="x", padx=15, pady=4)

        ttk.Label(frame_set, text="会議名:").grid(row=0, column=0, sticky="w", pady=2)
        self.entry_meeting = ttk.Entry(frame_set, width=44)
        self.entry_meeting.grid(row=0, column=1, padx=10, pady=2)

        ttk.Label(frame_set, text="対象画面:").grid(row=1, column=0, sticky="nw", pady=2)
        mon_frame = ttk.Frame(frame_set)
        mon_frame.grid(row=1, column=1, padx=10, pady=2, sticky="w")
        self.combo_monitor = ttk.Combobox(mon_frame, width=42, state="readonly")
        self.combo_monitor.pack(anchor="w")
        self.combo_monitor.bind("<<ComboboxSelected>>", self._on_monitor_changed)

        region_row = ttk.Frame(mon_frame)
        region_row.pack(anchor="w", pady=(3, 0))
        self.btn_region = ttk.Button(region_row, text="🖱 範囲を選択...",
            command=self._select_capture_region)
        self.btn_region.pack(side="left")
        self.btn_region_clear = ttk.Button(region_row, text="🗑 解除",
            width=7, command=self._clear_capture_region)
        self.btn_region_clear.pack(side="left", padx=(4, 0))
        self.label_region = ttk.Label(region_row, text="", font=("BIZ UDゴシック", 8),
            foreground="#6b7280")
        self.label_region.pack(side="left", padx=(8, 0))

        ttk.Label(frame_set, text="音声デバイス:").grid(row=2, column=0, sticky="nw", pady=2)
        audio_frame = ttk.Frame(frame_set)
        audio_frame.grid(row=2, column=1, padx=10, pady=2, sticky="w")

        lb_row = ttk.Frame(audio_frame)
        lb_row.pack(fill="x")
        self._audio_scroll = ttk.Scrollbar(lb_row, orient="vertical")
        self.listbox_audio = tk.Listbox(
            lb_row, selectmode="extended", height=5, width=52,
            font=("BIZ UDゴシック", 9), exportselection=False,
            yscrollcommand=self._audio_scroll.set)
        self._audio_scroll.config(command=self.listbox_audio.yview)
        self.listbox_audio.pack(side="left", fill="x", expand=True)
        self._audio_scroll.pack(side="right", fill="y")

        ttk.Label(audio_frame,
            text="※ Ctrl/Shift+クリックで複数選択（スピーカー・マイク混在可）",
            font=("BIZ UDゴシック", 8), foreground="#6b7280").pack(anchor="w", pady=(2, 0))
        audio_btn_row = ttk.Frame(audio_frame)
        audio_btn_row.pack(anchor="w", pady=(4, 0))
        self.btn_refresh_audio = ttk.Button(audio_btn_row, text="🔄 デバイス再検出",
            command=self._refresh_audio_devices)
        self.btn_refresh_audio.pack(side="left")
        self.btn_select_all_audio = ttk.Button(audio_btn_row, text="🎚 すべて選択",
            command=self._select_all_audio)
        self.btn_select_all_audio.pack(side="left", padx=(6, 0))

        ttk.Label(frame_set, text="動作モード:").grid(row=3, column=0, sticky="w", pady=2)
        self.combo_mode = ttk.Combobox(frame_set, width=42, state="readonly",
            values=["軽量（録音のみ）", "高負荷（リアルタイム文字起こし）"])
        self.combo_mode.grid(row=3, column=1, padx=10, pady=2)
        self.combo_mode.current(0)
        self.combo_mode.bind("<<ComboboxSelected>>", self._on_mode_changed)

        ttk.Label(frame_set, text="文字起こし言語:").grid(row=4, column=0, sticky="w", pady=2)
        self.combo_lang = ttk.Combobox(frame_set, width=42, state="readonly",
            values=[label for label, _ in LANG_OPTIONS])
        self.combo_lang.grid(row=4, column=1, padx=10, pady=2)
        self.combo_lang.current(0)
        self.combo_lang.bind("<<ComboboxSelected>>", self._on_lang_changed)

        frame_level = ttk.Frame(self.root, padding=(15, 2))
        frame_level.pack(fill="x")
        self._show_level_var = tk.BooleanVar(value=True)
        self.chk_level = tk.Checkbutton(
            frame_level, text="入力レベル:", variable=self._show_level_var,
            font=("BIZ UDゴシック", 9), command=self._toggle_level_meter)
        self.chk_level.pack(side="left")
        self.level_canvas = tk.Canvas(frame_level, height=14, bg="#1f2937",
            highlightthickness=1, highlightbackground="#374151")
        self.level_canvas.pack(side="left", fill="x", expand=True, padx=(8, 0))

        self._detect_monitors()

        frame_ctrl = ttk.Frame(self.root, padding=5)
        frame_ctrl.pack(fill="x", padx=15, pady=4)

        row_buttons = ttk.Frame(frame_ctrl)
        row_buttons.pack(fill="x", pady=3)
        self.btn_toggle = tk.Button(
            row_buttons, text="▶ 会議記録を開始", bg="#10b981", fg="white",
            font=("BIZ UDゴシック", 12, "bold"), relief="raised", padx=8, pady=8,
            command=self._toggle_recording)
        self.btn_toggle.pack(side="left", fill="x", expand=True)

        self.btn_manual_snap = tk.Button(
            row_buttons, text="📷 手動キャプチャ", bg="#6366f1", fg="white",
            font=("BIZ UDゴシック", 10, "bold"), relief="raised", padx=10, pady=8,
            command=self._manual_snapshot, state="disabled")
        self.btn_manual_snap.pack(side="right", padx=(6, 0))

        self.btn_marker = tk.Button(
            row_buttons, text="⭐ マーカー", bg="#f59e0b", fg="white",
            font=("BIZ UDゴシック", 10, "bold"), relief="raised", padx=10, pady=8,
            command=self._add_marker, state="disabled")
        self.btn_marker.pack(side="right", padx=(6, 0))

        row_post = ttk.Frame(frame_ctrl)
        row_post.pack(fill="x", pady=(3, 0))
        self.btn_postprocess = tk.Button(
            row_post, text="📄 議事録を生成（後処理）", bg="#0ea5e9", fg="white",
            font=("BIZ UDゴシック", 10, "bold"), relief="raised", padx=10, pady=6,
            command=self._run_postprocess)
        self.btn_postprocess.pack(side="left", fill="x", expand=True)
        self.btn_import_video = tk.Button(
            row_post, text="🎬 動画から生成", bg="#8b5cf6", fg="white",
            font=("BIZ UDゴシック", 10, "bold"), relief="raised", padx=10, pady=6,
            command=self._run_video_import)
        self.btn_import_video.pack(side="left", padx=(6, 0))
        self.btn_cancel_post = tk.Button(
            row_post, text="✖ 中断", bg="#ef4444", fg="white",
            font=("BIZ UDゴシック", 10, "bold"), relief="raised", padx=10, pady=6,
            command=self._cancel_postprocess, state="disabled")
        self.btn_cancel_post.pack(side="right", padx=(6, 0))

        self.var_video_snapshots = tk.BooleanVar(value=True)
        self.chk_video_snapshots = ttk.Checkbutton(
            frame_ctrl, text="動画入力時、画面変化を画像として保存・議事録に表示",
            variable=self.var_video_snapshots)
        self.chk_video_snapshots.pack(anchor="w", pady=(3, 0))

        row_prog = ttk.Frame(frame_ctrl)
        row_prog.pack(fill="x", pady=(4, 0))
        self.progress_post = ttk.Progressbar(row_prog, maximum=1000, value=0)
        self.progress_post.pack(side="left", fill="x", expand=True)
        self.label_progress = ttk.Label(row_prog, text="", width=30, anchor="w",
            font=("BIZ UDゴシック", 8), foreground="#6b7280")
        self.label_progress.pack(side="right", padx=(8, 0))

        self.label_status = ttk.Label(frame_ctrl, text="ステータス: 停止中",
            font=("BIZ UDゴシック", 10, "bold"), foreground="#6b7280")
        self.label_status.pack(anchor="w", pady=2)

        frame_log = ttk.LabelFrame(self.root, text=" 動作ログ ", padding=5)
        frame_log.pack(fill="both", expand=True, padx=15, pady=4)
        self.log_area = scrolledtext.ScrolledText(
            frame_log, height=7, font=("Consolas", 9), state="disabled")
        self.log_area.pack(fill="both", expand=True)

        frame_tr = ttk.LabelFrame(self.root, text=" 文字起こし（高負荷モード） ", padding=5)
        frame_tr.pack(fill="both", expand=True, padx=15, pady=(4, 8))
        self.transcript_area = scrolledtext.ScrolledText(
            frame_tr, height=6, font=("BIZ UDゴシック", 10), state="disabled")
        self.transcript_area.pack(fill="both", expand=True)
        self._set_transcript_enabled(False)

    # --------------------------------------------------------- Audio device enumeration

    # Host APIs that expose the same physical mic. Earlier = preferred.
    _HOSTAPI_PRIORITY = ("Windows WASAPI", "Windows WDM-KS", "Windows DirectSound", "MME")
    _HOSTAPI_SHORT = {
        "Windows WASAPI": "WASAPI", "Windows WDM-KS": "WDM-KS",
        "Windows DirectSound": "DSound", "MME": "MME",
    }

    # Virtual endpoints that just alias whatever the default input is. Listing
    # them means "select all" captures the default mic two or three times over.
    _VIRTUAL_INPUTS = (
        "サウンド マッパー", "サウンドマッパー", "sound mapper",
        "プライマリ サウンド キャプチャ", "primary sound capture",
    )

    @classmethod
    def _is_virtual_input(cls, name):
        low = name.lower()
        return any(v in low for v in cls._VIRTUAL_INPUTS)

    @staticmethod
    def _dedupe_key(name):
        """Normalized key matching one physical device across host APIs.
        MME truncates names to 31 chars, so compare a short normalized prefix."""
        return re.sub(r"[^0-9a-z぀-ヿ一-鿿]+", "", name.lower())[:30]

    def _enumerate_audio_devices(self):
        """Enumerate all capturable audio devices into a unified list.
        Each entry: {key, kind, native_idx, name, channels, rate, api}
          key: 'S<native_idx>' (speaker/loopback) / 'M<native_idx>' (mic)

        Mics are deduplicated across host APIs (a single jack is otherwise listed
        under MME / DirectSound / WASAPI / WDM-KS): selecting every entry would
        open the same hardware several times — some opens fail, and the ones that
        succeed get mixed into each other.
        """
        devices = []
        speaker_names = set()

        # --- Speaker/loopback via pyaudiowpatch ---
        p = pyaudio.PyAudio()
        try:
            for lb in p.get_loopback_device_info_generator():
                if lb.get("maxInputChannels", 0) > 0:
                    name = lb["name"]
                    speaker_names.add(name)
                    devices.append({
                        "key": f"S{lb['index']}",
                        "kind": "speaker",
                        "native_idx": lb["index"],
                        "name": name,
                        "channels": int(lb["maxInputChannels"]),
                        "rate": int(lb["defaultSampleRate"]),
                        "api": "WASAPI",
                    })
        except Exception as e:
            print(f"[デバイス列挙警告: speaker] {e}")
        finally:
            p.terminate()

        # --- Mic/input via sounddevice ---
        try:
            try:
                hostapis = list(sd.query_hostapis())
            except Exception:
                hostapis = []
            best = {}      # dedupe key -> (priority, entry)
            order = []     # dedupe keys in first-seen order
            skipped = 0
            for info in sd.query_devices():
                if info.get("max_input_channels", 0) <= 0:
                    continue
                name = info["name"]
                # Skip WASAPI loopbacks sounddevice also surfaces as inputs, and
                # any input already listed on the speaker side (dedupe by name).
                if "Loopback" in name or name in speaker_names:
                    continue
                if self._is_virtual_input(name):
                    continue
                api_idx = info.get("hostapi", -1)
                api_name = hostapis[api_idx]["name"] if 0 <= api_idx < len(hostapis) else ""
                try:
                    prio = self._HOSTAPI_PRIORITY.index(api_name)
                except ValueError:
                    prio = len(self._HOSTAPI_PRIORITY)
                entry = {
                    "key": f"M{info['index']}",
                    "kind": "mic",
                    "native_idx": info["index"],
                    "name": name,
                    "channels": int(info["max_input_channels"]),
                    "rate": int(info["default_samplerate"]),
                    "api": self._HOSTAPI_SHORT.get(api_name, api_name),
                }
                k = self._dedupe_key(name)
                if k not in best:
                    best[k] = (prio, entry)
                    order.append(k)
                else:
                    skipped += 1
                    if prio < best[k][0]:
                        best[k] = (prio, entry)
            devices.extend(best[k][1] for k in order)
            if skipped:
                print(f"[デバイス列挙] 別ホストAPIの重複マイク {skipped}件を非表示")
        except Exception as e:
            print(f"[デバイス列挙警告: mic] {e}")

        self._audio_devices = devices

    def _populate_audio_listbox(self):
        """Reflect self._audio_devices into the Listbox, then apply default selection."""
        if not getattr(self, "listbox_audio", None):
            return
        self.listbox_audio.delete(0, tk.END)
        for dev in self._audio_devices:
            prefix = "🔊 " if dev["kind"] == "speaker" else "🎤 "
            api = dev.get("api", "")
            suffix = f"  [{api}]" if api and api != "WASAPI" else ""
            self.listbox_audio.insert(tk.END, f"{prefix}{dev['name']}{suffix}")
        # Default selection: first speaker + first mic (mirrors prior "both on")
        first_speaker = first_mic = None
        for i, dev in enumerate(self._audio_devices):
            if dev["kind"] == "speaker" and first_speaker is None:
                first_speaker = i
            elif dev["kind"] == "mic" and first_mic is None:
                first_mic = i
        for idx in (first_speaker, first_mic):
            if idx is not None:
                self.listbox_audio.selection_set(idx)

    def _refresh_audio_devices(self):
        if self.is_recording:
            return  # devices are locked while recording
        self._enumerate_audio_devices()
        self._populate_audio_listbox()
        n_spk = sum(1 for d in self._audio_devices if d["kind"] == "speaker")
        n_mic = sum(1 for d in self._audio_devices if d["kind"] == "mic")
        self._log(f"オーディオデバイス再検出: スピーカー{n_spk} / マイク{n_mic}")

    def _select_all_audio(self):
        """Select every listed device. Dead/idle ones are tolerated at capture time."""
        if self.is_recording:
            return
        self.listbox_audio.selection_set(0, tk.END)
        self._log(f"音声デバイスを全選択: {self.listbox_audio.size()}件")

    def _selected_audio_devices(self):
        """Return list of device dicts currently selected in the Listbox."""
        sel = self.listbox_audio.curselection()
        return [self._audio_devices[i] for i in sel]

    @staticmethod
    def _audio_source_label_n(selected):
        """Backward-compatible AUDIO_SOURCE string for metadata.txt.
        1 speaker / 1 mic / 1+1 map to legacy values; otherwise a readable summary.
        """
        n_speaker = sum(1 for d in selected if d["kind"] == "speaker")
        n_mic = sum(1 for d in selected if d["kind"] == "mic")
        if n_speaker == 1 and n_mic == 0:
            return "speaker_loopback"
        if n_speaker == 0 and n_mic == 1:
            return "microphone"
        if n_speaker == 1 and n_mic == 1:
            return "both_mixed"
        parts = []
        if n_speaker:
            parts.append(f"speaker_loopback x{n_speaker}")
        if n_mic:
            parts.append(f"microphone x{n_mic}")
        return " + ".join(parts) if parts else "none"

    # ---------------------------------------------------- Level meter & timer

    def _toggle_level_meter(self):
        if self._show_level_var.get():
            self.level_canvas.pack(side="left", fill="x", expand=True, padx=(8, 0))
        else:
            self.level_canvas.pack_forget()

    def _tick_level_meter(self):
        if self._show_level_var.get():
            self.level_canvas.delete("all")
            w = self.level_canvas.winfo_width()
            if w > 1:
                level = self._audio_level
                bar_w = max(1, int(level * w))
                color = "#10b981" if level < 0.6 else "#f59e0b" if level < 0.85 else "#ef4444"
                self.level_canvas.create_rectangle(0, 0, bar_w, 14, fill=color, outline="")
        self.root.after(80, self._tick_level_meter)

    def _update_level(self, pcm_bytes):
        samples = np.frombuffer(pcm_bytes, dtype=np.int16)
        if len(samples) > 0:
            rms = np.sqrt(np.mean(samples.astype(np.float32) ** 2)) / 32768.0
            self._audio_level = min(1.0, rms * 3.0)

    def _tick_elapsed_timer(self):
        if self.is_recording:
            elapsed = time.time() - self._record_start
            m, s = divmod(int(elapsed), 60)
            h, m = divmod(m, 60)
            self.label_status.config(
                text=f"ステータス: 記録中  {h:02d}:{m:02d}:{s:02d}",
                foreground="#ef4444")
            self.root.after(500, self._tick_elapsed_timer)

    # ------------------------------------------------ Mode change & pre-load

    def _on_mode_changed(self, _=None):
        is_full = self.combo_mode.current() == 1
        self._set_transcript_enabled(is_full)
        if is_full and self._preloaded_transcriber is None:
            self._preload_model()

    def _preload_model(self):
        def _load():
            try:
                from faster_whisper import WhisperModel
                self.root.after(0, self._log,
                    f"文字起こしモデルを事前読み込み中... ({self.REALTIME_WHISPER_MODEL})")
                self._preloaded_transcriber = WhisperModel(
                    self.REALTIME_WHISPER_MODEL,
                    device=self.WHISPER_DEVICE, compute_type=self.WHISPER_COMPUTE)
                self.root.after(0, self._log,
                    f"文字起こしモデル読み込み完了 ({self.REALTIME_WHISPER_MODEL})")
            except Exception as e:
                self.root.after(0, self._log, f"[事前読み込みエラー] {e}")
        threading.Thread(target=_load, daemon=True).start()

    def _on_lang_changed(self, _=None):
        label = LANG_OPTIONS[self.combo_lang.current()][0]
        lang = LANG_OPTIONS[self.combo_lang.current()][1]
        # A single faster-whisper model handles ja/en/etc., so no model
        # reload is needed on language change — just update the target lang.
        self._rt_lang = None if lang == "auto" else lang
        self._log(f"文字起こし言語設定: {label}")

        if self.is_recording:
            self._write_language_event(lang, label)
        else:
            # Pre-load for next recording (no-op if already preloaded)
            if self.combo_mode.current() == 1 and self._preloaded_transcriber is None:
                self._preload_model()

    def _run_postprocess(self, folder=None):
        if self._post_running:
            messagebox.showinfo("実行中", "後処理がすでに実行中です。")
            return
        if folder is None:
            initial_dir = self._last_record_dir if self._last_record_dir else None
            folder = filedialog.askdirectory(
                title="議事録生成対象フォルダを選択", initialdir=initial_dir)
        if not folder:
            return
        self._post_running = True
        self._post_cancel.clear()
        self.btn_postprocess.config(state="disabled", text="⏳ 処理中...")
        self.btn_import_video.config(state="disabled")
        self.btn_cancel_post.config(state="normal", text="✖ 中断")
        self.progress_post.config(value=0)
        self.label_status.config(text="ステータス: 議事録生成中...", foreground="#f59e0b")
        threading.Thread(target=self._postprocess_worker, args=(folder,), daemon=True).start()

    def _run_video_import(self):
        """Select a video, extract its audio, then run the normal post-process."""
        if self._post_running:
            messagebox.showinfo("実行中", "後処理がすでに実行中です。")
            return
        video_path = filedialog.askopenfilename(
            title="議事録を生成する動画を選択",
            filetypes=[
                ("動画ファイル", "*.mp4 *.mov *.mkv *.webm *.avi *.m4v *.mts *.m2ts"),
                ("すべてのファイル", "*.*"),
            ])
        if not video_path:
            return
        if not os.path.exists(FFMPEG_PATH):
            messagebox.showerror("ffmpegが見つかりません",
                f"動画の音声抽出にffmpeg.exeが必要です。\n\n期待パス:\n{FFMPEG_PATH}")
            return

        lang = LANG_OPTIONS[self.combo_lang.current()][1]
        capture_scenes = self.var_video_snapshots.get()
        self._post_running = True
        self._post_cancel.clear()
        self.btn_postprocess.config(state="disabled")
        self.btn_import_video.config(state="disabled", text="⏳ 取込中...")
        self.btn_cancel_post.config(state="normal", text="✖ 中断")
        self.progress_post.config(value=0)
        self.label_status.config(text="ステータス: 動画を取込中...", foreground="#8b5cf6")
        threading.Thread(target=self._video_import_worker,
                         args=(video_path, lang, capture_scenes), daemon=True).start()

    def _video_import_worker(self, video_path, lang, capture_scenes):
        result = "失敗"
        try:
            folder = import_video_file(
                video_path, language=lang, progress=self._post_progress,
                cancel=self._post_cancel.is_set, capture_scenes=capture_scenes,
                scene_interval=self.INTERVAL,
                scene_threshold=self.DHASH_THRESHOLD,
                max_edge=self.CAPTURE_MAX_EDGE,
                jpeg_quality=self.JPEG_QUALITY)
            self._last_record_dir = folder
            post_process_folder(folder, progress=self._post_progress,
                                cancel=self._post_cancel.is_set)
            result = "成功"
        except PostProcessCancelled:
            result = "中断"
        except Exception as e:
            self.root.after(0, self._log, f"[動画取込エラー] {e}")
            traceback.print_exc()
        finally:
            self._post_running = False
            self.root.after(0, self._log, f"動画からの議事録生成完了 ({result})")
            self.root.after(0, self._finish_postprocess_ui, result)

    def _cancel_postprocess(self):
        """Ask the worker to stop; it checks the flag at each phase boundary."""
        self._post_cancel.set()
        self.btn_cancel_post.config(state="disabled", text="中断中...")
        self._log("後処理の中断を要求しました（現在の処理が終わり次第停止します）")

    def _post_progress(self, frac, msg):
        """Progress callback — called from the worker thread."""
        self.root.after(0, self._apply_post_progress, frac, msg)

    def _apply_post_progress(self, frac, msg):
        if frac is not None:
            self.progress_post.config(value=max(0, min(1000, int(frac * 1000))))
        short = msg if len(msg) <= 30 else msg[:29] + "…"
        self.label_progress.config(text=short)
        self._log(f"[後処理] {msg}")

    def _postprocess_worker(self, folder):
        result = "失敗"
        try:
            post_process_folder(folder, progress=self._post_progress,
                                cancel=self._post_cancel.is_set)
            result = "成功"
        except PostProcessCancelled:
            result = "中断"
        except Exception as e:
            self.root.after(0, self._log, f"[後処理エラー] {e}")
            traceback.print_exc()
        finally:
            self._post_running = False
            self.root.after(0, self._log, f"後処理完了 ({result})")
            self.root.after(0, self._finish_postprocess_ui, result)

    def _finish_postprocess_ui(self, result):
        self.btn_postprocess.config(state="normal", text="📄 議事録を生成（後処理）")
        self.btn_import_video.config(state="normal", text="🎬 動画から生成")
        self.btn_cancel_post.config(state="disabled", text="✖ 中断")
        self.label_progress.config(text="" if result == "成功" else result)
        if result != "成功":
            self.progress_post.config(value=0)
        if not self.is_recording:
            self.label_status.config(text="ステータス: 停止中", foreground="#6b7280")

    # ------------------------------------------------------- Capture region

    @staticmethod
    def _display_scale(root, sct):
        """Physical pixels per tkinter unit.

        1.0 for a DPI-unaware process (Windows virtualizes both tkinter and the
        screen grab identically), but a DPI-aware host would make tkinter report
        logical units while mss keeps reporting physical ones.
        """
        try:
            logical = root.winfo_screenwidth()
            physical = sct.monitors[1]["width"] if len(sct.monitors) > 1 else logical
            if logical > 0 and physical > 0:
                return physical / logical
        except Exception as e:
            print(f"[DPI取得警告] {e}")
        return 1.0

    def _is_region_mode(self):
        return self.combo_monitor.get() == REGION_CHOICE

    def _update_region_label(self):
        region = self.CAPTURE_REGION
        if region:
            self.label_region.config(
                text=f"{region['width']}×{region['height']} "
                     f"(x={region['left']}, y={region['top']})",
                foreground="#0ea5e9")
        else:
            self.label_region.config(text="未選択", foreground="#6b7280")
        on = self._is_region_mode() and not self.is_recording
        self.btn_region.config(state="normal" if on else "disabled")
        self.btn_region_clear.config(
            state="normal" if on and region else "disabled")

    def _on_monitor_changed(self, _=None):
        self._update_region_label()
        if self._is_region_mode() and not self.CAPTURE_REGION:
            self._select_capture_region()

    def _clear_capture_region(self):
        self.CAPTURE_REGION = None
        self._save_settings()
        self._update_region_label()
        self._log("キャプチャ範囲を解除しました")

    def _select_capture_region(self):
        """Freeze the screen, dim it, and let the user drag out a rectangle.

        The overlay shows a still screenshot rather than being see-through: a
        genuinely transparent window would be click-through on Windows, which
        would swallow the drag we are trying to capture.
        """
        if self.is_recording:
            return
        try:
            with mss.MSS() as sct:
                virt = dict(sct.monitors[0])
                scale = self._display_scale(self.root, sct)
                shot = sct.grab(virt)
            base = Image.frombytes("RGB", shot.size, shot.bgra, "raw", "BGRX")
        except Exception as e:
            messagebox.showerror("エラー", f"画面を取得できません:\n{e}")
            return

        win_w = max(1, int(round(virt["width"] / scale)))
        win_h = max(1, int(round(virt["height"] / scale)))
        win_x = int(round(virt["left"] / scale))
        win_y = int(round(virt["top"] / scale))
        if (win_w, win_h) != base.size:
            base = base.resize((win_w, win_h), Image.LANCZOS)

        win = tk.Toplevel(self.root)
        win.overrideredirect(True)
        win.attributes("-topmost", True)
        win.geometry(f"{win_w}x{win_h}+{win_x}+{win_y}")
        canvas = tk.Canvas(win, width=win_w, height=win_h, highlightthickness=0,
                           cursor="crosshair", bg="black")
        canvas.pack()
        photo = ImageTk.PhotoImage(base)
        canvas.create_image(0, 0, image=photo, anchor="nw")
        canvas.image = photo   # keep a reference or Tk drops the image

        # Four dimmed panels around the selection give a real "hole" without
        # any window transparency, and updating them is four coordinate sets.
        shades = [canvas.create_rectangle(0, 0, 0, 0, fill="black",
                                          stipple="gray75", outline="")
                  for _ in range(4)]
        box = canvas.create_rectangle(0, 0, 0, 0, outline="#38bdf8", width=2)
        size_text = canvas.create_text(0, 0, text="", anchor="nw", fill="#ffffff",
                                       font=("BIZ UDゴシック", 12, "bold"))
        canvas.create_text(win_w // 2, 20, anchor="n", fill="#ffffff",
                           font=("BIZ UDゴシック", 13, "bold"),
                           text="ドラッグして範囲を選択    /    Esc でキャンセル")

        def _shade(x0, y0, x1, y1):
            canvas.coords(shades[0], 0, 0, win_w, y0)          # above
            canvas.coords(shades[1], 0, y1, win_w, win_h)      # below
            canvas.coords(shades[2], 0, y0, x0, y1)            # left
            canvas.coords(shades[3], x1, y0, win_w, y1)        # right

        # Before the first drag there is no hole to leave, and the four-panel
        # split cannot cover the screen on its own — stretch one panel over it.
        canvas.coords(shades[0], 0, 0, win_w, win_h)
        state = {"x0": 0, "y0": 0, "dragging": False, "rect": None}

        def _corners(event):
            return (min(state["x0"], event.x), min(state["y0"], event.y),
                    max(state["x0"], event.x), max(state["y0"], event.y))

        def _on_press(event):
            state.update(x0=event.x, y0=event.y, dragging=True)

        def _on_move(event):
            if not state["dragging"]:
                return
            x0, y0, x1, y1 = _corners(event)
            canvas.coords(box, x0, y0, x1, y1)
            _shade(x0, y0, x1, y1)
            canvas.itemconfig(size_text, text=f"{int((x1 - x0) * scale)} × "
                                              f"{int((y1 - y0) * scale)}")
            canvas.coords(size_text, x0 + 6, y0 + 6 if y0 + 30 < win_h else y0 - 26)

        def _on_release(event):
            if not state["dragging"]:
                return
            state["dragging"] = False
            x0, y0, x1, y1 = _corners(event)
            state["rect"] = {
                "left": int(round(virt["left"] + x0 * scale)),
                "top": int(round(virt["top"] + y0 * scale)),
                "width": int(round((x1 - x0) * scale)),
                "height": int(round((y1 - y0) * scale)),
            }
            win.destroy()

        canvas.bind("<ButtonPress-1>", _on_press)
        canvas.bind("<B1-Motion>", _on_move)
        canvas.bind("<ButtonRelease-1>", _on_release)
        # A borderless topmost window has no close button, so cancelling must
        # not hinge on one binding landing: Escape from anywhere in the app,
        # and right-click on the overlay itself.
        cancel = lambda _e=None: win.destroy()
        for widget in (win, canvas):
            widget.bind("<Escape>", cancel)
            widget.bind("<Button-3>", cancel)
        win.bind_all("<Escape>", cancel)

        def _poll_escape():
            """Watch the physical Escape key, not just Tk's focused widget.

            A borderless topmost window can end up without keyboard focus — if
            that happens the key bindings never fire and the overlay would be
            unclosable, so ask Windows directly instead.
            """
            if not win.winfo_exists():
                return
            try:
                import ctypes
                if ctypes.windll.user32.GetAsyncKeyState(0x1B) & 0x8000:  # VK_ESCAPE
                    win.destroy()
                    return
            except Exception as e:
                print(f"[範囲選択警告] Escape 監視を停止: {e}")
                return
            win.after(60, _poll_escape)

        if os.name == "nt":
            win.after(60, _poll_escape)
        win.focus_force()
        canvas.focus_set()
        win.grab_set()
        try:
            self.root.wait_window(win)
        finally:
            try:
                win.unbind_all("<Escape>")
            except tk.TclError:
                pass

        rect = state["rect"]
        if not rect:
            self._log("範囲選択をキャンセルしました")
            self._update_region_label()
            return
        if rect["width"] < MIN_REGION or rect["height"] < MIN_REGION:
            messagebox.showwarning("範囲が小さすぎます",
                f"{MIN_REGION}×{MIN_REGION} ピクセル以上を選択してください。\n"
                f"（選択: {rect['width']}×{rect['height']}）")
            self._update_region_label()
            return

        self.CAPTURE_REGION = rect
        if not self._is_region_mode():
            values = list(self.combo_monitor["values"])
            if REGION_CHOICE in values:
                self.combo_monitor.current(values.index(REGION_CHOICE))
        self._save_settings()
        self._update_region_label()
        self._log(f"キャプチャ範囲: {rect['width']}×{rect['height']} "
                  f"(x={rect['left']}, y={rect['top']})")

    def _capture_rect(self, sct):
        """The rectangle to grab — a chosen region, or the selected display."""
        if self._is_region_mode() and self.CAPTURE_REGION:
            return dict(self.CAPTURE_REGION)
        idx = min(max(self._recording_mon_idx, 0), len(sct.monitors) - 1)
        return dict(sct.monitors[idx])

    def _grab_image(self, sct, rect):
        """Grab a frame and shrink it if its long edge exceeds the limit."""
        shot = sct.grab(rect)
        img = Image.frombytes("RGB", shot.size, shot.bgra, "raw", "BGRX")
        limit = self.CAPTURE_MAX_EDGE
        if limit and max(img.size) > limit:
            img.thumbnail((limit, limit), Image.LANCZOS)
        return img

    def _validate_region(self):
        """Reject a stale region — displays get unplugged and resolutions change."""
        if not (self._is_region_mode() and self.CAPTURE_REGION):
            return True
        region = self.CAPTURE_REGION
        with mss.MSS() as sct:
            virt = sct.monitors[0]
        inside = (region["left"] >= virt["left"]
                  and region["top"] >= virt["top"]
                  and region["left"] + region["width"] <= virt["left"] + virt["width"]
                  and region["top"] + region["height"] <= virt["top"] + virt["height"])
        if inside:
            return True
        return messagebox.askyesno(
            "範囲が画面外です",
            f"保存されている範囲 ({region['left']},{region['top']} "
            f"{region['width']}×{region['height']}) が現在の画面に収まりません。\n"
            "ディスプレイ構成が変わった可能性があります。\n\n"
            "このまま録画を開始しますか？（画面外は黒く記録されます）")

    # ------------------------------------------------------ Recordings browser

    def _scan_recordings(self):
        """List recording folders under the app directory and the year folders.

        Scanned lazily when the window opens, not at startup — a few hundred
        meetings would otherwise slow every launch.
        """
        roots = [BASE_DIR]
        try:
            for name in os.listdir(BASE_DIR):
                path = os.path.join(BASE_DIR, name)
                if os.path.isdir(path) and re.fullmatch(r"\d{6,8}", name):
                    roots.append(path)   # e.g. an IC-recorder import folder
        except OSError as e:
            print(f"[記録一覧警告] {e}")

        seen = set()
        rows = []
        for root in roots:
            try:
                names = sorted(os.listdir(root), reverse=True)
            except OSError:
                continue
            for name in names:
                path = os.path.join(root, name)
                if path in seen or not os.path.isdir(path):
                    continue
                if not any(os.path.exists(os.path.join(path, a))
                           for a in ("audio_main.mp3", "audio_main.wav")):
                    continue
                seen.add(path)
                meta = _read_metadata(path)
                images = sum(1 for f in os.listdir(path)
                             if f.lower().endswith((".jpg", ".jpeg", ".png")))
                rows.append({
                    "path": path,
                    "name": meta.get("MEETING_NAME") or name,
                    "start": meta.get("START_TIME_STR", ""),
                    "images": images,
                    "markers": len(_read_jsonl(os.path.join(path, "markers.jsonl"))),
                    "roles": bool(meta.get("ROLE_TRACKS")),
                    "report": os.path.exists(os.path.join(path, "meeting_report.md")),
                    "html": os.path.exists(os.path.join(path, "meeting_report.html")),
                })
        rows.sort(key=lambda r: (r["start"] or "", r["path"]), reverse=True)
        return rows

    def _open_recordings_browser(self):
        if getattr(self, "_browser_win", None) is not None:
            try:
                self._browser_win.lift()
                return
            except tk.TclError:
                pass

        win = tk.Toplevel(self.root)
        self._browser_win = win
        win.title("GijirokuStudio - 記録一覧")
        win.geometry("900x520")
        win.transient(self.root)

        top = ttk.Frame(win, padding=(10, 8))
        top.pack(fill="x")
        ttk.Label(top, text="検索:").pack(side="left")
        var_q = tk.StringVar()
        entry_q = ttk.Entry(top, width=30, textvariable=var_q)
        entry_q.pack(side="left", padx=(6, 10))
        var_full = tk.BooleanVar(value=False)
        ttk.Checkbutton(top, text="議事録の本文も検索", variable=var_full).pack(side="left")
        lbl_count = ttk.Label(top, text="", foreground="#6b7280")
        lbl_count.pack(side="right")

        cols = ("start", "name", "images", "markers", "roles", "report")
        tree = ttk.Treeview(win, columns=cols, show="headings", selectmode="browse")
        for col, text, width in (
                ("start", "開始時刻", 150), ("name", "会議名", 300),
                ("images", "画像", 60), ("markers", "⭐", 50),
                ("roles", "話者分離", 80), ("report", "議事録", 150)):
            tree.heading(col, text=text)
            tree.column(col, width=width, anchor="w")
        scroll = ttk.Scrollbar(win, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scroll.set)
        tree.pack(side="left", fill="both", expand=True, padx=(10, 0), pady=4)
        scroll.pack(side="left", fill="y", padx=(0, 10), pady=4)

        rows = []
        by_item = {}

        def _matches(row, query, full_text):
            if not query:
                return True
            low = query.lower()
            if low in row["name"].lower() or low in row["start"].lower():
                return True
            if not full_text:
                return False
            for fname in ("meeting_report.md", "transcription.txt"):
                path = os.path.join(row["path"], fname)
                if not os.path.exists(path):
                    continue
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        if low in f.read().lower():
                            return True
                except OSError:
                    pass
            return False

        def _refill(*_):
            query = var_q.get().strip()
            tree.delete(*tree.get_children())
            by_item.clear()
            shown = 0
            for row in rows:
                if not _matches(row, query, var_full.get()):
                    continue
                report = ("md + html" if row["html"] else "md") if row["report"] else "未生成"
                item = tree.insert("", "end", values=(
                    row["start"] or "-", row["name"], row["images"],
                    row["markers"] or "", "✓" if row["roles"] else "",
                    report))
                by_item[item] = row
                shown += 1
            lbl_count.config(text=f"{shown} / {len(rows)} 件")

        def _reload():
            rows.clear()
            rows.extend(self._scan_recordings())
            _refill()

        def _selected():
            sel = tree.selection()
            return by_item.get(sel[0]) if sel else None

        def _open_folder():
            row = _selected()
            if row:
                os.startfile(row["path"])

        def _open_report(kind):
            row = _selected()
            if not row:
                return
            path = os.path.join(row["path"], f"meeting_report.{kind}")
            if os.path.exists(path):
                os.startfile(path)
            else:
                messagebox.showinfo("未生成",
                    f"meeting_report.{kind} がありません。先に後処理を実行してください。")

        def _run_post():
            row = _selected()
            if not row:
                return
            win.destroy()
            self._browser_win = None
            self._run_postprocess(row["path"])

        bar = ttk.Frame(win, padding=(10, 8))
        bar.pack(side="bottom", fill="x")
        for text, cmd in (("📄 議事録を生成/再生成", _run_post),
                          ("📝 Markdown", lambda: _open_report("md")),
                          ("🌐 HTML", lambda: _open_report("html")),
                          ("📁 フォルダを開く", _open_folder),
                          ("🔄 再スキャン", _reload)):
            ttk.Button(bar, text=text, command=cmd).pack(side="left", padx=(0, 6))

        # Trace the variable rather than <KeyRelease> so paste, clear, and
        # programmatic edits all refresh the list.
        var_q.trace_add("write", _refill)
        var_full.trace_add("write", _refill)
        tree.bind("<Double-1>", lambda _e: _run_post())

        def _on_close():
            self._browser_win = None
            win.destroy()

        win.protocol("WM_DELETE_WINDOW", _on_close)
        _reload()

    def _open_settings_dialog(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("GijirokuStudio - 設定")
        dlg.resizable(False, False)
        dlg.transient(self.root)
        dlg.grab_set()

        frame = ttk.LabelFrame(dlg, text=" 詳細設定 ", padding=15)
        frame.pack(padx=15, pady=15)

        # dHash threshold
        ttk.Label(frame, text="差分感度（dHash閾値）:").grid(row=0, column=0, sticky="w", pady=4)
        var_dhash = tk.IntVar(value=self.DHASH_THRESHOLD)
        scale_dhash = tk.Scale(frame, from_=1, to=30, orient="horizontal",
            variable=var_dhash, length=220, showvalue=False)
        scale_dhash.grid(row=0, column=1, padx=(10, 4), pady=4)
        lbl_dhash = ttk.Label(frame, text=str(self.DHASH_THRESHOLD), width=4)
        lbl_dhash.grid(row=0, column=2, pady=4)
        var_dhash.trace_add("write", lambda *_: lbl_dhash.config(text=str(var_dhash.get())))

        # Audio gain
        ttk.Label(frame, text="音声ゲイン:").grid(row=1, column=0, sticky="w", pady=4)
        var_gain = tk.DoubleVar(value=self.AUDIO_GAIN)
        scale_gain = tk.Scale(frame, from_=1.0, to=5.0, resolution=0.1,
            orient="horizontal", variable=var_gain, length=220, showvalue=False)
        scale_gain.grid(row=1, column=1, padx=(10, 4), pady=4)
        lbl_gain = ttk.Label(frame, text=f"{self.AUDIO_GAIN:.1f}", width=4)
        lbl_gain.grid(row=1, column=2, pady=4)
        var_gain.trace_add("write", lambda *_: lbl_gain.config(text=f"{var_gain.get():.1f}"))

        # JPEG quality
        ttk.Label(frame, text="JPEG画質:").grid(row=2, column=0, sticky="w", pady=4)
        var_jpeg = tk.IntVar(value=self.JPEG_QUALITY)
        scale_jpeg = tk.Scale(frame, from_=50, to=100, orient="horizontal",
            variable=var_jpeg, length=220, showvalue=False)
        scale_jpeg.grid(row=2, column=1, padx=(10, 4), pady=4)
        lbl_jpeg = ttk.Label(frame, text=str(self.JPEG_QUALITY), width=4)
        lbl_jpeg.grid(row=2, column=2, pady=4)
        var_jpeg.trace_add("write", lambda *_: lbl_jpeg.config(text=str(var_jpeg.get())))

        # Capture long-edge limit
        ttk.Label(frame, text="画像の長辺上限:").grid(row=3, column=0, sticky="w", pady=4)
        var_edge = tk.IntVar(value=self.CAPTURE_MAX_EDGE)
        scale_edge = tk.Scale(frame, from_=0, to=3840, resolution=160,
            orient="horizontal", variable=var_edge, length=220, showvalue=False)
        scale_edge.grid(row=3, column=1, padx=(10, 4), pady=4)
        lbl_edge = ttk.Label(frame, text=str(self.CAPTURE_MAX_EDGE), width=5)
        lbl_edge.grid(row=3, column=2, pady=4)
        var_edge.trace_add("write", lambda *_: lbl_edge.config(
            text=(str(var_edge.get()) if var_edge.get() else "無制限")))
        ttk.Label(frame, text="※ 0 で無効。超えた場合だけ縮小します",
            font=("BIZ UDゴシック", 8), foreground="#6b7280").grid(
            row=4, column=0, columnspan=3, sticky="w")

        # Slide OCR
        var_ocr = tk.BooleanVar(value=self.OCR_ENABLED)
        ttk.Checkbutton(frame, text="スライドOCR（後処理で画像から文字を抽出）",
            variable=var_ocr).grid(row=5, column=0, columnspan=3, sticky="w",
                                   pady=(8, 0))
        ttk.Label(frame, text="※ Windows の日本語OCR言語パックが必要です",
            font=("BIZ UDゴシック", 8), foreground="#6b7280").grid(
            row=6, column=0, columnspan=3, sticky="w")

        # AI summary
        sum_frame = ttk.LabelFrame(dlg, text=" AI要約（後処理） ", padding=15)
        sum_frame.pack(padx=15, pady=(0, 10), fill="x")

        ttk.Label(sum_frame, text="プロバイダ:").grid(row=0, column=0, sticky="w", pady=4)
        current = next((i for i, (_, v) in enumerate(SUMMARY_PROVIDERS)
                        if v == self.SUMMARY_PROVIDER), 0)
        combo_prov = ttk.Combobox(sum_frame, width=34, state="readonly",
            values=[label for label, _ in SUMMARY_PROVIDERS])
        combo_prov.grid(row=0, column=1, padx=(10, 0), pady=4)
        combo_prov.current(current)

        ttk.Label(sum_frame, text="モデル:").grid(row=1, column=0, sticky="w", pady=4)
        entry_model = ttk.Entry(sum_frame, width=36)
        entry_model.grid(row=1, column=1, padx=(10, 0), pady=4)
        entry_model.insert(0, self.SUMMARY_MODEL)

        ttk.Label(sum_frame,
            text="※ 空欄で既定（Claude: claude-opus-5 / Ollama: qwen3）\n"
                 "※ Claude API を選ぶと文字起こしが外部に送信されます",
            font=("BIZ UDゴシック", 8), foreground="#b45309",
            justify="left").grid(row=2, column=0, columnspan=2, sticky="w", pady=(4, 0))

        # Buttons
        btn_frame = ttk.Frame(dlg)
        btn_frame.pack(padx=15, pady=(0, 15))

        def _ok():
            self.DHASH_THRESHOLD = var_dhash.get()
            self.AUDIO_GAIN = var_gain.get()
            self.JPEG_QUALITY = var_jpeg.get()
            self.CAPTURE_MAX_EDGE = var_edge.get()
            self.OCR_ENABLED = var_ocr.get()
            self.SUMMARY_PROVIDER = SUMMARY_PROVIDERS[combo_prov.current()][1]
            self.SUMMARY_MODEL = entry_model.get().strip()
            self._save_settings()
            self._log(f"設定更新: dHash={self.DHASH_THRESHOLD}, ゲイン={self.AUDIO_GAIN:.1f}, "
                      f"JPEG={self.JPEG_QUALITY}, OCR={'有効' if self.OCR_ENABLED else '無効'}, "
                      f"要約={self.SUMMARY_PROVIDER}")
            dlg.destroy()

        def _cancel():
            dlg.destroy()

        tk.Button(btn_frame, text="  OK  ", command=_ok,
            font=("BIZ UDゴシック", 10)).pack(side="left", padx=8)
        tk.Button(btn_frame, text="キャンセル", command=_cancel,
            font=("BIZ UDゴシック", 10)).pack(side="left", padx=8)

        dlg.protocol("WM_DELETE_WINDOW", _cancel)

        # Center dialog on parent
        dlg.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - dlg.winfo_width()) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - dlg.winfo_height()) // 2
        dlg.geometry(f"+{x}+{y}")

    def _reset_settings(self):
        self.DHASH_THRESHOLD = 10
        self.AUDIO_GAIN = 2.0
        self.JPEG_QUALITY = 85
        self._save_settings()
        self._log("設定をデフォルトにリセットしました")

    def _set_transcript_enabled(self, on):
        self.transcript_area.config(
            state="normal" if on else "disabled",
            background="#fffef0" if on else "#f3f4f6")

    # -------------------------------------------------------- Language segment log

    def _write_language_event(self, lang, label):
        if not self._recording_dir:
            return
        elapsed = time.time() - self._record_start
        path = os.path.join(self._recording_dir, "language_segments.jsonl")
        try:
            with open(path, "a", encoding="utf-8") as f:
                f.write(json.dumps({
                    "elapsed": round(elapsed, 3),
                    "lang": lang,
                    "label": label
                }, ensure_ascii=False) + "\n")
        except Exception as e:
            print(f"[言語セグメント書き込み警告] {e}")

    def _start_hotkey(self):
        """Register the marker hotkey for the duration of the recording."""
        self._stop_hotkey()
        if not self.MARKER_HOTKEY:
            return
        self._hotkey = _GlobalHotkey(self.MARKER_HOTKEY, self._add_marker,
                                     log=lambda m: self.root.after(0, self._log, m))
        if self._hotkey.start():
            self._log(f"マーカーのホットキー: {self.MARKER_HOTKEY}（他アプリ使用中でも有効）")

    def _stop_hotkey(self):
        if self._hotkey is not None:
            self._hotkey.stop()
            self._hotkey = None

    def _add_marker(self, label=None):
        """Bookmark the current moment. Callable from the UI or the hotkey thread."""
        if not self.is_recording or not self._recording_dir:
            return
        elapsed = time.time() - self._record_start
        entry = {
            "elapsed": round(elapsed, 3),
            "epoch": round(time.time(), 3),
            "label": label or "重要",
        }
        try:
            with self._marker_lock:
                path = os.path.join(self._recording_dir, "markers.jsonl")
                with open(path, "a", encoding="utf-8") as f:
                    f.write(json.dumps(entry, ensure_ascii=False) + "\n")
                self._marker_count += 1
            count = self._marker_count
        except Exception as e:
            self.root.after(0, self._log, f"[マーカー書き込み警告] {e}")
            return
        m, s = divmod(int(elapsed), 60)
        # May arrive from the hotkey thread, so touch the UI via after().
        self.root.after(0, self._log,
                        f"⭐ マーカー {count}: [{m:02d}:{s:02d}] {entry['label']}")

    def _write_snapshot_log(self, fname, snap_type, diff):
        if not self._recording_dir:
            return
        elapsed = time.time() - self._record_start
        epoch = time.time()
        path = os.path.join(self._recording_dir, "snapshots.jsonl")
        try:
            with open(path, "a", encoding="utf-8") as f:
                f.write(json.dumps({
                    "file": fname,
                    "epoch": round(epoch, 3),
                    "elapsed": round(elapsed, 3),
                    "type": snap_type,
                    "diff": int(diff) if diff is not None else None,
                }, ensure_ascii=False) + "\n")
        except Exception as e:
            # A dropped line means the image never reaches the report, so this
            # has to be visible in the app — a console print is not.
            self.root.after(0, self._log, f"[スナップショット記録エラー] {fname}: {e}")

    # ------------------------------------------------------------ Logging

    def _manual_snapshot(self):
        if not self.is_recording or self._recording_dir is None:
            return
        try:
            with mss.MSS() as sct:
                img = self._grab_image(sct, self._recording_rect or
                                       self._capture_rect(sct))
            now = datetime.datetime.now()
            ts = now.strftime('%H%M%S')
            ms = f"{now.microsecond // 1000:03d}"
            fname = f"manual_{ts}_{ms}.jpg"
            img.save(os.path.join(self._recording_dir, fname), "JPEG", quality=self.JPEG_QUALITY)
            self._log(f"手動キャプチャ -> {fname}")
            self._write_snapshot_log(fname, "manual", None)
        except Exception as e:
            self._log(f"[手動キャプチャエラー] {e}")

    def _log(self, text):
        self.log_area.config(state="normal")
        self.log_area.insert(tk.END,
            f"[{datetime.datetime.now().strftime('%H:%M:%S')}] {text}\n")
        self.log_area.see(tk.END)
        self.log_area.config(state="disabled")

    def _log_transcript(self, text):
        self.transcript_area.config(state="normal")
        ts = datetime.datetime.now().strftime('%H:%M:%S')
        self.transcript_area.insert(tk.END, f"[{ts}] {text}\n")
        self.transcript_area.see(tk.END)
        self.transcript_area.config(state="disabled")

    def _detect_monitors(self):
        with mss.MSS() as sct:
            vals = []
            for i, m in enumerate(sct.monitors):
                if i == 0:
                    vals.append(f"画面 [0]: 全画面 (Virtual Screen)")
                else:
                    vals.append(f"画面 [{i}]: ディスプレイ {i} ({m['width']}x{m['height']})")
            vals.append(REGION_CHOICE)
            self.combo_monitor["values"] = vals
            self.combo_monitor.current(0)

    # -------------------------------------------------------------- Control

    @staticmethod
    def _sanitize_name(name):
        """Make a string safe to embed in a folder/file name (Windows)."""
        cleaned = re.sub(r'[\\/:*?"<>|]', "_", name)   # invalid filename chars
        cleaned = re.sub(r"\s+", " ", cleaned).strip()  # collapse whitespace
        cleaned = cleaned.rstrip(". ")                  # trailing dots/spaces
        return cleaned

    def _toggle_recording(self):
        if not self.is_recording:
            if not os.path.exists(FFMPEG_PATH):
                self._log("[エラー] ffmpeg.exe が見つかりません")
                messagebox.showerror("エラー",
                    f"ffmpeg.exe が見つかりません。\n\n期待パス:\n{FFMPEG_PATH}")
                return
            selected = self._selected_audio_devices()
            if not selected:
                messagebox.showwarning("音声デバイス未選択",
                    "録音するオーディオデバイスを1つ以上選択してください。")
                return
            if not self._validate_region():
                return
            self._meeting_name = self.entry_meeting.get().strip()
            self.is_recording = True
            self.stop_event.clear()
            self._pcm_written = 0
            self._queue_overflow_count = 0
            self._record_start = time.time()
            self._recording_mon_idx = self.combo_monitor.current()
            with mss.MSS() as sct:
                self._recording_rect = self._capture_rect(sct)
            self._marker_count = 0
            self.btn_toggle.config(text="■ 記録を停止して保存", bg="#ef4444")
            self.btn_manual_snap.config(state="normal")
            self.btn_marker.config(state="normal")
            self._start_hotkey()
            for w in (self.combo_monitor, self.combo_mode):
                w.config(state="disabled")
            self.listbox_audio.config(state="disabled")
            self.btn_refresh_audio.config(state="disabled")
            self.btn_select_all_audio.config(state="disabled")
            self.btn_region.config(state="disabled")
            self.btn_region_clear.config(state="disabled")
            self.entry_meeting.config(state="disabled")
            # combo_lang stays enabled for mid-recording language switching
            self._tick_elapsed_timer()
            self._pipeline_thread = threading.Thread(target=self._pipeline, daemon=True)
            self._pipeline_thread.start()
        else:
            # UI resets instantly; cleanup runs in background
            self.is_recording = False
            self.stop_event.set()
            self._stop_hotkey()
            self.btn_toggle.config(state="disabled")
            self.btn_manual_snap.config(state="disabled")
            self.btn_marker.config(state="disabled")
            self.label_status.config(text="ステータス: 保存処理中...", foreground="#f59e0b")
            threading.Thread(target=self._async_cleanup, daemon=True).start()

    def _reset_ui(self):
        self.is_recording = False
        self.stop_event.clear()
        self._audio_level = 0.0
        self._queue_overflow_count = 0
        self._recording_dir = None
        self._stop_hotkey()
        self.btn_toggle.config(text="▶ 会議記録を開始", bg="#10b981", state="normal")
        self.btn_manual_snap.config(state="disabled")
        self.btn_marker.config(state="disabled")
        self.label_status.config(text="ステータス: 停止中", foreground="#6b7280")
        for w in (self.combo_monitor, self.combo_mode):
            w.config(state="readonly")
        self.listbox_audio.config(state="normal")
        self.btn_refresh_audio.config(state="normal")
        self.btn_select_all_audio.config(state="normal")
        self._recording_rect = None
        self._update_region_label()
        self.entry_meeting.config(state="normal")

    # ----------------------------------------------------------- Main pipeline

    def _pipeline(self):
        now = datetime.datetime.now()
        stamp = now.strftime("%Y%m%d_%H%M%S")
        safe = self._sanitize_name(self._meeting_name)
        suffix = safe if safe else "Meeting"
        dir_name = os.path.join(BASE_DIR, f"{stamp}_{suffix}")
        os.makedirs(dir_name, exist_ok=True)
        self._recording_dir = dir_name
        self.root.after(0, self._log, f"フォルダー作成: {dir_name}")

        lang = LANG_OPTIONS[self.combo_lang.current()][1]
        label = LANG_OPTIONS[self.combo_lang.current()][0]
        self._write_language_event(lang, label)

        start_epoch = time.time()
        selected = self._selected_audio_devices()
        n = len(selected)
        mode_full = self.combo_mode.current() == 1
        out_mp3 = os.path.join(dir_name, "audio_main.mp3")

        # One dedicated queue per selected device
        self._audio_queues = [queue.Queue(maxsize=200) for _ in range(n)]
        self.transcribe_queue = queue.Queue(maxsize=400) if mode_full else None

        # Speaker separation needs both sides captured; with one side there is
        # nothing to separate and the role tracks would just duplicate the mix.
        n_mic = sum(1 for d in selected if d["kind"] == "mic")
        n_spk = sum(1 for d in selected if d["kind"] == "speaker")
        separate = n > 1 and n_mic >= 1 and n_spk >= 1

        if not self._start_writers(dir_name, out_mp3, separate):
            self.root.after(0, self._log, "[エラー] ffmpeg 起動失敗")
            self.root.after(0, self._reset_ui)
            return

        if mode_full:
            self._transcription_start(dir_name)

        active = threading.Event(); active.set()
        threads = []
        for i, dev in enumerate(selected):
            threads.append(threading.Thread(
                target=self._capture_device,
                args=(active, dev, self._audio_queues[i]), daemon=True))
        for t in threads:
            t.start()

        if n == 1:
            writer_t = threading.Thread(
                target=self._single_writer,
                args=(active, self._audio_queues[0]), daemon=True)
            writer_t.start()
            mixer_t = None
        else:
            mixer_t = threading.Thread(
                target=self._mixer_loop_n,
                args=(active, self._audio_queues, selected), daemon=True)
            mixer_t.start()
            writer_t = None

        # Screen capture
        next_target = time.time() + self.INTERVAL
        last_hash = None
        snap_count = 0
        with mss.MSS() as sct:
            monitor = self._recording_rect or self._capture_rect(sct)
            self.root.after(0, self._log,
                f"キャプチャ対象: {monitor['width']}×{monitor['height']} "
                f"(x={monitor['left']}, y={monitor['top']})")
            while self.is_recording and not self.stop_event.is_set():
                dt = next_target - time.time()
                if dt > 0:
                    time.sleep(dt)
                try:
                    img = self._grab_image(sct, monitor)
                    h = imagehash.dhash(img)
                except Exception as e:
                    self.root.after(0, self._log, f"[画面エラー] {e}")
                    next_target += self.INTERVAL
                    continue

                diff = 0
                changed = last_hash is None or (h - last_hash) > self.DHASH_THRESHOLD
                if changed:
                    if last_hash is not None:
                        # int(): imagehash returns numpy.int64, which json
                        # cannot serialize — that silently dropped every
                        # snapshot after the first one from the log.
                        diff = int(h - last_hash)
                    snap_count += 1
                    ts = datetime.datetime.now().strftime('%H%M%S')
                    ms = f"{datetime.datetime.now().microsecond // 1000:03d}"
                    fname = f"snapshot_{ts}_{ms}.jpg"
                    img.save(os.path.join(dir_name, fname), "JPEG", quality=self.JPEG_QUALITY)
                    self.root.after(0, self._log, f"画面変化検知 -> {fname} (diff={diff})")
                    self._write_snapshot_log(fname, "auto", diff)
                    last_hash = h
                next_target += self.INTERVAL

        # Save context for async cleanup
        self._pipeline_ctx = {
            'active': active, 'threads': threads,
            'mixer_t': mixer_t, 'writer_t': writer_t,
            'mode_full': mode_full, 'start_epoch': start_epoch,
            'now': now, 'dir_name': dir_name,
            'audio_src_str': self._audio_source_label_n(selected),
            'snap_count': snap_count, 'lang': lang,
            'meeting_name': self._meeting_name,
            'role_tracks': sorted(self._role_writers),
        }

    def _async_cleanup(self):
        """Run teardown in background so UI stays responsive."""
        # Wait for pipeline thread to finish setting _pipeline_ctx
        if self._pipeline_thread:
            self._pipeline_thread.join(timeout=10)
            self._pipeline_thread = None
        ctx = self._pipeline_ctx
        if ctx is None:
            self._stop_writers()
            self.root.after(0, self._reset_ui)
            return

        active = ctx['active']
        active.clear()
        for t in ctx['threads']:
            t.join(timeout=3)
        if ctx['mixer_t']:
            ctx['mixer_t'].join(timeout=5)
        if ctx['writer_t']:
            ctx['writer_t'].join(timeout=3)

        self._stop_writers()
        if ctx['mode_full']:
            self._transcription_stop()

        dir_name = ctx['dir_name']
        with open(os.path.join(dir_name, "metadata.txt"), "w", encoding="utf-8") as f:
            f.write(f"MEETING_NAME={ctx['meeting_name']}\n")
            f.write(f"START_TIME_EPOCH={ctx['start_epoch']}\n")
            f.write(f"START_TIME_STR={ctx['now'].strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"AUDIO_SOURCE={ctx['audio_src_str']}\n")
            f.write(f"MODE={'full' if ctx['mode_full'] else 'light'}\n")
            f.write(f"SNAPSHOT_COUNT={ctx['snap_count']}\n")
            f.write(f"MARKER_COUNT={self._marker_count}\n")
            f.write("AUDIO_FILE=audio_main.mp3\n")
            f.write(f"LANGUAGE={ctx['lang']}\n")
            if ctx['role_tracks']:
                f.write(f"ROLE_TRACKS={','.join(ctx['role_tracks'])}\n")
                f.write(f"AUDIO_SELF_FILE={ROLE_TRACK_SELF}\n")
                f.write(f"AUDIO_OTHER_FILE={ROLE_TRACK_OTHER}\n")
            if ctx['mode_full']:
                f.write("TRANSCRIPTION_FILE=transcription.txt\n")

        elapsed = time.time() - ctx['start_epoch']
        m, s = divmod(int(elapsed), 60)
        audio_sec = self._pcm_written / (self.TARGET_RATE * 4)   # stereo int16
        msg = (f"記録完了 ({m:02d}:{s:02d})。画像: {ctx['snap_count']}枚 "
               f"/ 音声: {audio_sec:.0f}秒 / 保存先: {dir_name}")
        if self._queue_overflow_count > 0:
            msg += f" / ⚠ キュー溢れ: {self._queue_overflow_count}回"
        self._pipeline_ctx = None
        self._last_record_dir = dir_name
        self.root.after(0, self._log, msg)
        if self._pcm_written == 0:
            self.root.after(0, messagebox.showwarning, "音声なし",
                "音声が1バイトも記録されませんでした。\n\n"
                "選択したデバイスがすべて無音／使用不可の可能性があります。\n"
                "動作ログの「[スキップ]」「無音（データ未着）」を確認してください。\n\n"
                f"フォルダー:\n{dir_name}")
        else:
            self.root.after(0, messagebox.showinfo, "完了",
                f"すべての記録が正常に保存されました。\n\n"
                f"音声: 約{audio_sec:.0f}秒\n\nフォルダー:\n{dir_name}")
        self.root.after(0, self._reset_ui)

    # --------------------------------------------------- FFmpeg (Captura pattern)

    def _start_writers(self, dir_name, out_mp3, separate):
        """Open the mixed-audio encoder and, when separating speakers, the role
        tracks. Returns False if the main encoder could not start."""
        self.ffmpeg_proc = None
        self._role_writers = {}
        try:
            self.ffmpeg_proc = _PcmWriter(out_mp3, self.TARGET_RATE)
        except Exception as e:
            self.root.after(0, self._log, f"[ffmpegエラー] {e}")
            return False
        if not separate:
            return True
        for role, fname in ((ROLE_SELF, ROLE_TRACK_SELF),
                            (ROLE_OTHER, ROLE_TRACK_OTHER)):
            path = os.path.join(dir_name, fname)
            try:
                self._role_writers[role] = _PcmWriter(
                    path, self.TARGET_RATE, transcription_only=True)
            except Exception as e:
                self.root.after(0, self._log, f"[話者トラック警告] {role}: {e}")
        if self._role_writers:
            self.root.after(0, self._log,
                "話者分離を有効化（自分=マイク / 相手=スピーカー）")
        return True

    def _write_audio(self, main_pcm, role_chunks=None):
        """Write one output chunk to the mix and to every role track.

        Each role track advances by exactly as many bytes as the main mix —
        silence where that role contributed nothing — so the three files stay
        on one timeline and transcript timestamps remain comparable.
        """
        self._pcm_written += len(main_pcm)
        if self.ffmpeg_proc is not None:
            self.ffmpeg_proc.write(main_pcm)
        if not self._role_writers:
            return
        silence = None
        for role, writer in self._role_writers.items():
            chunk = role_chunks.get(role) if role_chunks else None
            if chunk is None:
                if silence is None:
                    silence = bytes(len(main_pcm))
                chunk = silence
            writer.write(chunk)

    def _stop_writers(self):
        for writer in list(self._role_writers.values()):
            writer.close()
        self._role_writers = {}
        if self.ffmpeg_proc is not None:
            self.ffmpeg_proc.close()
            self.ffmpeg_proc = None
            self.root.after(0, self._log, "ffmpeg エンコード完了 -> MP3 保存済み")

    # ---------------------------------------------------- Audio normalization

    def _to_stereo_s16le(self, raw_bytes, channels, rate):
        channels = max(1, int(channels))
        # Keep whole frames only: a short/odd tail would make reshape() raise
        # inside the audio callback.
        usable = len(raw_bytes) - (len(raw_bytes) % (2 * channels))
        if usable <= 0:
            return b""
        data = np.frombuffer(raw_bytes, dtype=np.int16,
                             count=usable // 2).astype(np.float32)
        if channels == 1:
            data = np.column_stack([data, data]).flatten()
        elif channels > 2:
            data = data.reshape(-1, channels)[:, :2].flatten()
        if rate != self.TARGET_RATE and len(data) >= 4:
            stereo = data.reshape(-1, 2)
            n_in = len(stereo)
            n_out = max(1, int(n_in * self.TARGET_RATE / rate))
            x_in = np.arange(n_in, dtype=np.float64)
            x_out = np.linspace(0, n_in - 1, n_out)
            left = np.interp(x_out, x_in, stereo[:, 0])
            right = np.interp(x_out, x_in, stereo[:, 1])
            data = np.column_stack([left, right]).flatten()
        data *= self.AUDIO_GAIN
        np.clip(data, -32768, 32767, out=data)
        return data.astype(np.int16).tobytes()

    # ------------------------------------------------- Transcribe forwarding

    def _push_transcribe(self, pcm):
        """Forward a PCM chunk to the transcription queue (no-op in light mode)."""
        if self.transcribe_queue is not None:
            try:
                self.transcribe_queue.put_nowait(pcm)
            except queue.Full:
                self._queue_overflow_count += 1

    # ------------------------------------------- Single-source writer thread

    def _single_writer(self, active, src_queue):
        """Direct writer for the single-device case (no mixing)."""
        while (active.is_set() or not src_queue.empty()) and not self.stop_event.is_set():
            try:
                pcm = src_queue.get(timeout=0.2)
            except queue.Empty:
                continue
            if pcm is _EOF:
                break
            self._update_level(pcm)
            self._write_audio(pcm)
            self._push_transcribe(pcm)

    # ------------------------------------------------- Mixer thread (N inputs)

    def _mixer_loop_n(self, active, queues, devices=None):
        """Mix N capture queues into one PCM stream.

        Each input queue may receive the _EOF sentinel to mark end-of-stream for
        that source. Sources that reach EOF contribute silence; mixing continues
        with the remaining live sources until every source is exhausted.

        A source may also stop delivering without ever reaching EOF: Windows does
        not run the audio engine for a loopback endpoint nothing is playing to,
        so an idle speaker produces no callbacks at all. Such a source is treated
        as silent after STARVE_SECONDS instead of blocking the whole mix.
        """
        N = len(queues)
        MIX_FRAMES = 1024
        FRAME_BYTES = 4                      # stereo int16 = 4 bytes
        MIX_BYTES = MIX_FRAMES * FRAME_BYTES  # 4096 bytes per output chunk
        OVERFLOW = MIX_BYTES * 50            # ~1.2 sec buffer limit
        BUF_CAP = MIX_BYTES * 400            # ~9 sec hard cap, memory backstop

        bufs = [bytearray() for _ in range(N)]
        eof = [False] * N
        t_start = time.time()
        last_data = [t_start] * N
        starved = [False] * N
        stop_deadline = None
        emitted = 0                 # bytes pushed downstream, for silence padding
        SILENCE = bytes(MIX_BYTES)

        # Which speaker each source belongs to: mics are me, loopbacks are them.
        # Left as None when no role tracks are open, so the sub-mix work below
        # is skipped entirely rather than computed and thrown away.
        role_of = [None] * N
        if devices and self._role_writers:
            role_of = [ROLE_SELF if d.get("kind") == "mic" else ROLE_OTHER
                       for d in devices[:N]] + [None] * max(0, N - len(devices))

        def _name(i):
            if devices and i < len(devices):
                return devices[i]["name"]
            return f"source{i}"

        def _emit(raw, role_chunks=None):
            nonlocal emitted
            emitted += len(raw)
            self._update_level(raw)
            self._write_audio(raw, role_chunks)
            self._push_transcribe(raw)

        def _solo(i):
            """A single source's chunk, tagged with its role for the role track."""
            raw = bytes(bufs[i][:MIX_BYTES])
            del bufs[i][:MIX_BYTES]
            return raw, ({role_of[i]: raw} if role_of[i] else None)

        def _drain(i):
            for _ in range(30):
                item = self._qget(queues[i], timeout=0)
                if item is None:
                    return
                if item is _EOF:
                    eof[i] = True
                    return
                bufs[i].extend(item)
                last_data[i] = time.time()

        while not self.stop_event.is_set():
            live = [i for i in range(N) if not eof[i]]

            got = False
            for i in live:
                before = len(bufs[i])
                _drain(i)
                if len(bufs[i]) != before or eof[i]:
                    got = True

            now = time.time()
            for i in live:
                quiet = now - last_data[i] >= self.STARVE_SECONDS
                if quiet != starved[i]:
                    starved[i] = quiet
                    state = "無音（データ未着）" if quiet else "受信再開"
                    self.root.after(0, self._log, f"[音声] {_name(i)}: {state}")

            # Mix aligned chunks. Sources that are live but have gone quiet do not
            # hold the mix back — they simply contribute nothing to these chunks.
            while True:
                ready = [i for i in live if len(bufs[i]) >= MIX_BYTES]
                blocking = [i for i in live
                            if len(bufs[i]) < MIX_BYTES and not starved[i]]
                if not ready or blocking:
                    break
                acc = np.zeros(MIX_BYTES // 2, dtype=np.float32)
                by_role = {}
                for i in ready:
                    chunk = np.frombuffer(bytes(bufs[i][:MIX_BYTES]),
                        dtype=np.int16).astype(np.float32)
                    acc += chunk
                    if role_of[i] is not None:
                        by_role.setdefault(role_of[i], []).append(chunk)
                    del bufs[i][:MIX_BYTES]
                acc *= (1.0 / len(ready))    # average-mix: amplitude stays ≤1.0 as sources grow
                np.clip(acc, -32768, 32767, out=acc)
                # Each role track averages over its OWN sources only — dividing
                # by the global count would quietly halve one speaker's volume.
                role_chunks = {}
                for role, chunks in by_role.items():
                    racc = np.sum(chunks, axis=0)
                    racc *= (1.0 / len(chunks))
                    np.clip(racc, -32768, 32767, out=racc)
                    role_chunks[role] = racc.astype(np.int16).tobytes()
                _emit(acc.astype(np.int16).tobytes(), role_chunks)

            # Overflow fallback: one source bloated while others lag → emit solo
            for i in live:
                others_short = all(
                    len(bufs[k]) < MIX_BYTES
                    for k in range(N) if k != i and not eof[k])
                if others_short and len(bufs[i]) > OVERFLOW:
                    while len(bufs[i]) >= MIX_BYTES:
                        _emit(*_solo(i))

            # Keep the timeline honest. With only idle loopback devices selected
            # nothing arrives at all, and the MP3 would end up shorter than the
            # meeting — transcript timestamps would no longer line up with the
            # snapshots. Pad the gap with real silence, but only while every live
            # source is starved, so normal mixing is never second-guessed.
            if live and all(starved[i] for i in live) and not self.stop_event.is_set():
                deficit = int((now - t_start) * self.TARGET_RATE * FRAME_BYTES) - emitted
                while deficit >= MIX_BYTES:
                    _emit(SILENCE)
                    deficit -= MIX_BYTES

            # Memory backstop: never let a stalled mix grow without bound
            for i in range(N):
                if len(bufs[i]) > BUF_CAP:
                    drop = len(bufs[i]) - BUF_CAP
                    drop -= drop % FRAME_BYTES
                    del bufs[i][:drop]
                    self._queue_overflow_count += 1

            # Termination: stop requested and every source reached EOF. A capture
            # thread wedged in the driver must not hold the file open forever.
            if not active.is_set():
                if all(eof[i] for i in range(N)):
                    break
                if stop_deadline is None:
                    stop_deadline = now + 3.0
                elif now > stop_deadline:
                    self.root.after(0, self._log,
                        "[音声] 応答しないデバイスを待たずにミキシングを終了")
                    break

            if not got:
                time.sleep(0.01)

        # Flush remaining whole-frame data
        for i in range(N):
            usable = len(bufs[i]) - (len(bufs[i]) % FRAME_BYTES)
            if usable:
                raw = bytes(bufs[i][:usable])
                _emit(raw, {role_of[i]: raw} if role_of[i] else None)

    @staticmethod
    def _qget(q, timeout=0.05):
        if q is None:
            return None
        try:
            return q.get(timeout=timeout)
        except queue.Empty:
            return None

    # --------------------------------------------------- Audio capture (generic)

    def _capture_device(self, active, dev, out_queue):
        """Capture one device; normalize PCM into out_queue. Emits _EOF on exit."""
        own_com = _com_initialize()
        try:
            if dev["kind"] == "speaker":
                self._capture_speaker_dev(active, dev, out_queue)
            else:
                self._capture_mic_dev(active, dev, out_queue)
        except Exception as e:
            self.root.after(0, self._log, f"[キャプチャエラー] {dev['name']}: {e}")
        finally:
            if own_com:
                _com_uninitialize()   # after the stream is closed, never before
            # Guarantee the EOF sentinel lands so the mixer can drain & terminate.
            for _ in range(3):
                try:
                    out_queue.put_nowait(_EOF)
                    break
                except queue.Full:
                    try:
                        out_queue.get_nowait()  # drop one stale PCM to make room
                    except queue.Empty:
                        break

    # ----------------------------------------- Speaker capture (pyaudiowpatch)

    @staticmethod
    def _format_candidates(channels, rate):
        """(channels, rate) combos to try, best first. A device's advertised
        default is not always openable — multi-channel endpoints in particular."""
        combos = []
        for c in (channels, 2, 1):
            c = int(c)
            if c < 1:
                continue
            for r in (rate, 48000, 44100):
                r = int(r)
                if (c, r) not in combos:
                    combos.append((c, r))
        return combos

    def _capture_speaker_dev(self, active, dev, out_queue):
        def _make_callback(ch, sr):
            def _callback(in_data, frame_count, time_info, status):
                # An exception raised here propagates into PortAudio's C callback
                # and can take the process down — never let one escape.
                try:
                    pcm = self._to_stereo_s16le(in_data, ch, sr)
                    if pcm:
                        out_queue.put_nowait(pcm)
                except queue.Full:
                    pass
                except Exception:
                    pass
                return (None, pyaudio.paContinue)
            return _callback

        stream = None
        open_err = None   # first failure = the device's own advertised format
        # PortAudio must be initialised on the very thread that opens the stream:
        # its WASAPI backend sets up COM per thread, so a handle created on
        # another thread makes every open fail with "invalid sample rate".
        with self._open_lock:
            p = pyaudio.PyAudio()
            for ch, sr in self._format_candidates(dev["channels"], dev["rate"]):
                try:
                    stream = p.open(
                        format=pyaudio.paInt16, channels=ch, rate=sr,
                        input=True, input_device_index=dev["native_idx"],
                        frames_per_buffer=1024, start=False,
                        stream_callback=_make_callback(ch, sr))
                    stream.start_stream()
                    self.root.after(0, self._log,
                        f"スピーカーキャプチャ開始: {dev['name']} "
                        f"({sr}Hz, {ch}ch, idx={dev['native_idx']})")
                    break
                except Exception as e:
                    if open_err is None:
                        open_err = e
                    if stream is not None:
                        try:
                            stream.close()
                        except Exception:
                            pass
                        stream = None

        def _terminate():
            with self._open_lock:
                try:
                    p.terminate()
                except Exception as e:
                    print(f"[PyAudio終了警告] {dev['name']}: {e}")

        if stream is None:
            # Device unusable (disabled, exclusive-mode, unsupported format) —
            # report it and let the rest of the selection keep recording.
            _terminate()
            self.root.after(0, self._log,
                f"[スキップ] スピーカー {dev['name']}: 開けません ({open_err})")
            return

        try:
            while active.is_set() and self.is_recording:
                time.sleep(0.5)
        finally:
            try:
                stream.stop_stream()
            except Exception as e:
                print(f"[停止警告] {dev['name']}: {e}")
            try:
                stream.close()
            except Exception as e:
                print(f"[クローズ警告] {dev['name']}: {e}")
            _terminate()

    # ------------------------------------------- Mic capture (sounddevice)

    def _capture_mic_dev(self, active, dev, out_queue):
        def _make_callback(ch, sr):
            def _callback(in_data, frames, time_info, status):
                try:
                    pcm = self._to_stereo_s16le(in_data.tobytes(), ch, sr)
                    if pcm:
                        out_queue.put_nowait(pcm)
                except queue.Full:
                    pass
                except Exception:
                    pass
            return _callback

        stream = None
        open_err = None   # first failure = the device's own advertised format
        with self._open_lock:
            for ch, sr in self._format_candidates(min(dev["channels"], 2), dev["rate"]):
                try:
                    stream = sd.InputStream(
                        device=dev["native_idx"], channels=ch, samplerate=sr,
                        dtype='int16', blocksize=1024, callback=_make_callback(ch, sr))
                    stream.start()
                    self.root.after(0, self._log,
                        f"マイクキャプチャ開始: {dev['name']} ({sr}Hz, {ch}ch)")
                    break
                except Exception as e:
                    if open_err is None:
                        open_err = e
                    if stream is not None:
                        try:
                            stream.close()
                        except Exception:
                            pass
                        stream = None

        if stream is None:
            self.root.after(0, self._log,
                f"[スキップ] マイク {dev['name']}: 開けません ({open_err})")
            return

        try:
            while active.is_set() and self.is_recording:
                time.sleep(0.5)
        finally:
            try:
                stream.stop()
            except Exception as e:
                print(f"[停止警告] {dev['name']}: {e}")
            try:
                stream.close()
            except Exception as e:
                print(f"[クローズ警告] {dev['name']}: {e}")

    # ------------------------------------------ Transcription (faster-whisper)

    def _transcription_start(self, dir_name):
        try:
            if self._preloaded_transcriber is not None:
                self.transcriber = self._preloaded_transcriber
                self._preloaded_transcriber = None
                self.root.after(0, self._log, "事前読み込み済みモデルを使用")
            else:
                from faster_whisper import WhisperModel
                self.root.after(0, self._log,
                    f"文字起こしモデルを読み込み中... ({self.REALTIME_WHISPER_MODEL})")
                self.transcriber = WhisperModel(
                    self.REALTIME_WHISPER_MODEL,
                    device=self.WHISPER_DEVICE, compute_type=self.WHISPER_COMPUTE)

            lang = LANG_OPTIONS[self.combo_lang.current()][1]
            self._rt_lang = None if lang == "auto" else lang

            self.transcription_file = open(
                os.path.join(dir_name, "transcription.txt"), "w", encoding="utf-8")
            self._transcribe_start = time.time()
            self.root.after(0, self._log, "文字起こしスレッド起動")
            threading.Thread(target=self._transcribe_loop, daemon=True).start()

        except Exception as e:
            self.root.after(0, self._log, f"[文字起こしエラー] {e}")
            self.transcriber = None

    def _transcribe_loop(self):
        """Chunk-driven real-time transcription loop.

        Incoming PCM (44100Hz stereo int16, from self.transcribe_queue) is
        down-mixed to mono and resampled to 16000Hz. Once ~5 seconds of audio
        has accumulated, faster-whisper transcribes the chunk (auto-detecting
        ja/en when self._rt_lang is None). ~1 second of trailing audio is kept
        as overlap so words aren't cut at chunk boundaries.
        """
        from scipy.signal import resample as _resample

        chunk_samples = int(self.TRANSCRIBE_RATE * self.TRANSCRIBE_CHUNK_SECONDS)
        overlap_samples = int(self.TRANSCRIBE_RATE * self.TRANSCRIBE_OVERLAP_SECONDS)
        min_flush_samples = int(self.TRANSCRIBE_RATE * 0.5)

        buf = np.zeros(0, dtype=np.float32)
        last_lang = None

        def _drop_backlog():
            # Backpressure: if the queue is backing up, drop the oldest
            # chunks so transcription keeps pace with real time.
            q = self.transcribe_queue
            if q is None:
                return
            qsize = q.qsize()
            if qsize > 200:
                for _ in range(qsize - 50):
                    try:
                        q.get_nowait()
                        self._queue_overflow_count += 1
                    except queue.Empty:
                        break

        while self.is_recording and not self.stop_event.is_set():
            if self.transcribe_queue is None or self.transcriber is None:
                time.sleep(0.1)
                continue

            _drop_backlog()

            try:
                pcm = self.transcribe_queue.get(timeout=1.0)
            except queue.Empty:
                continue

            try:
                samples = np.frombuffer(pcm, dtype=np.int16).astype(np.float32) / 32768.0
                mono_44k = samples.reshape(-1, 2).mean(axis=1) if len(samples) >= 2 else samples
                n_out = int(len(mono_44k) * self.TRANSCRIBE_RATE / self.TARGET_RATE)
                if n_out > 0:
                    mono_16k = _resample(mono_44k, n_out).astype(np.float32)
                    buf = np.concatenate([buf, mono_16k])

                while len(buf) >= chunk_samples:
                    clip = buf[:chunk_samples]
                    last_lang = self._transcribe_chunk(clip, last_lang)
                    buf = buf[max(0, chunk_samples - overlap_samples):]
            except Exception as e:
                self.root.after(0, self._log, f"[文字起こし処理エラー] {e}")

        if len(buf) >= min_flush_samples and self.transcriber is not None:
            try:
                self._transcribe_chunk(buf, last_lang)
            except Exception as e:
                self.root.after(0, self._log, f"[バッファフラッシュ警告] {e}")

    def _transcribe_chunk(self, clip, prev_lang):
        """Transcribe one ~5s 16kHz mono chunk; return the language used."""
        try:
            lang = detect_ja_en(self.transcriber, clip) if self._rt_lang is None else self._rt_lang
            segments, info = self.transcriber.transcribe(
                clip, language=lang, vad_filter=True, beam_size=1)
            for seg in segments:
                text = seg.text.strip()
                if not text:
                    continue
                self.root.after(0, self._log_transcript, text)
                if self.transcription_file:
                    elapsed = time.time() - self._transcribe_start
                    self.transcription_file.write(f"[{elapsed:.1f}s] {text}\n")
                    self.transcription_file.flush()
            if self._rt_lang is None and lang != prev_lang:
                label = "日本語" if lang == "ja" else "English"
                self._write_language_event(lang, label)
            return lang
        except Exception as e:
            self.root.after(0, self._log, f"[文字起こし処理エラー] {e}")
            return prev_lang

    def _transcription_stop(self):
        # WhisperModel needs no explicit close/stop — just drop the reference.
        self.transcriber = None
        if self.transcription_file:
            self.transcription_file.close()
            self.transcription_file = None
            self.root.after(0, self._log, "文字起こし完了 -> transcription.txt")
        self.transcribe_queue = None


# ======================================================================
# CLI: Post-processing — whisper transcription + markdown generation
# ======================================================================

def load_app_settings():
    """Read settings.json into a dict ({} if absent/invalid). Shared by GUI & CLI."""
    if not os.path.exists(SETTINGS_PATH):
        return {}
    try:
        with open(SETTINGS_PATH, "r", encoding="utf-8") as f:
            s = json.load(f)
        return s if isinstance(s, dict) else {}
    except Exception as e:
        print(f"[設定読み込み警告] {e}")
        return {}


def load_glossary():
    """Load the term glossary CSV. Each row: form, reading, alias1, alias2, ...
    Returns list of {form, reading, aliases}. Empty list if the file is absent.
    Blank lines and '#'-prefixed lines are skipped.
    """
    if not os.path.exists(GLOSSARY_PATH):
        return []
    import csv
    terms = []
    try:
        with open(GLOSSARY_PATH, "r", encoding="utf-8-sig", newline="") as f:
            for row in csv.reader(f):
                row = [c.strip() for c in row]
                if not row or not row[0] or row[0].startswith("#"):
                    continue
                if row[0].lower() in ("正規形", "form", "用語", "term", "word", "name"):
                    continue  # header row
                form = row[0]
                reading = row[1] if len(row) > 1 else ""
                aliases = [a for a in row[2:] if a]
                terms.append({"form": form, "reading": reading, "aliases": aliases})
    except Exception as e:
        print(f"[用語辞書読み込み警告] {e}")
    return terms


def build_whisper_prompt(terms):
    """Build a Whisper initial_prompt (vocabulary bias) from canonical forms.
    Caps at ~80 terms / 400 chars to stay within the prompt token budget.
    Returns None when there is nothing to bias with.
    """
    forms = [t["form"] for t in terms if t.get("form")]
    if not forms:
        return None
    return "、".join(forms[:80])[:400]


def detect_ja_en(model, audio_f32):
    """Detect whether a 16kHz mono float32 clip is Japanese or English.

    Uses faster-whisper's model.detect_language(), restricted to just the
    ja/en pair. Falls back to a transcribe()-based detection (reading
    info.language, coercing anything that isn't 'en' to 'ja') for older
    faster-whisper versions that lack detect_language(). Defaults to 'ja'
    if detection fails entirely.
    """
    try:
        _, _, all_language_probs = model.detect_language(audio_f32)
        probs = dict(all_language_probs)
        ja_p = probs.get("ja", 0.0)
        en_p = probs.get("en", 0.0)
        return "en" if en_p > ja_p else "ja"
    except Exception:
        pass
    try:
        _, info = model.transcribe(audio_f32, language=None)
        return "en" if getattr(info, "language", "ja") == "en" else "ja"
    except Exception:
        return "ja"


def iter_speech_windows(audio_f32, sr, window_s=25.0):
    """Yield (start_sample, end_sample) windows covering speech in audio_f32.

    Uses faster-whisper's bundled Silero VAD (get_speech_timestamps) to find
    speech spans, then greedily groups consecutive spans into windows up to
    ~window_s seconds each. On any failure, or if no speech is detected,
    falls back to fixed contiguous windows of window_s seconds spanning the
    whole array. Windows shorter than ~0.5s are never yielded.
    """
    min_samples = int(0.5 * sr)
    window_samples = int(window_s * sr)
    total = len(audio_f32)

    def _fixed_windows():
        start = 0
        while start < total:
            end = min(start + window_samples, total)
            if end - start >= min_samples:
                yield (start, end)
            start = end

    try:
        from faster_whisper.vad import get_speech_timestamps
        spans = get_speech_timestamps(audio_f32)
        if not spans:
            yield from _fixed_windows()
            return

        cur_start = spans[0]["start"]
        cur_end = spans[0]["end"]
        for sp in spans[1:]:
            if sp["end"] - cur_start <= window_samples:
                cur_end = sp["end"]
            else:
                if cur_end - cur_start >= min_samples:
                    yield (cur_start, cur_end)
                cur_start = sp["start"]
                cur_end = sp["end"]
        if cur_end - cur_start >= min_samples:
            yield (cur_start, cur_end)
    except Exception:
        yield from _fixed_windows()


def apply_glossary(lines, terms):
    """Exact longest-match replacement of registered aliases -> canonical form.
    Correctly-recognised canonical forms are left untouched. Mutates `lines`
    in place; returns the number of lines that were changed.
    """
    repl = {}
    for t in terms:
        form = t.get("form")
        if not form:
            continue
        for pat in t.get("aliases", []):
            if pat and pat != form:
                repl[pat] = form
    if not repl:
        return 0
    patterns = sorted(repl.keys(), key=len, reverse=True)  # longest first
    changed = 0
    for ln in lines:
        text = ln.get("text", "")
        for pat in patterns:
            if pat in text:
                text = text.replace(pat, repl[pat])
        if text != ln.get("text", ""):
            ln["text"] = text
            changed += 1
    return changed


def decode_audio_16k_mono(path):
    """Decode any audio file to a float32 mono 16kHz array (whisper's input)."""
    import wave
    tmp = f"{path}._tmp16k.wav"
    flags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
    subprocess.run(
        [FFMPEG_PATH, "-y", "-i", path, "-ar", "16000", "-ac", "1", "-f", "wav", tmp],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True,
        creationflags=flags)
    try:
        with wave.open(tmp, "rb") as wf:
            sr = wf.getframerate()
            raw = wf.readframes(wf.getnframes())
    finally:
        try:
            os.remove(tmp)
        except OSError:
            pass
    return np.frombuffer(raw, dtype=np.int16).astype(np.float32) / 32768.0, sr


class PostProcessCancelled(Exception):
    """Raised at a checkpoint when the caller asked to stop post-processing."""


class _Reporter:
    """Progress reporting plus cancellation, threaded through post-processing.

    Calling it announces a phase; `check()` is the cheap cancel-only probe for
    tight loops where a log line per iteration would be noise.
    """

    def __init__(self, progress=None, cancel=None):
        self._progress = progress
        self._cancel = cancel

    def __call__(self, frac, msg):
        print(msg)
        if self._progress is not None:
            try:
                self._progress(frac, msg)
            except Exception as e:
                print(f"[進捗通知警告] {e}")
        self.check()

    def check(self):
        if self._cancel is not None and self._cancel():
            raise PostProcessCancelled()


# ------------------------------------------------------- Meeting data loading

def _read_jsonl(path):
    """Read a .jsonl log, tolerating a truncated or corrupt trailing line."""
    rows = []
    if not os.path.exists(path):
        return rows
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                rows.append(json.loads(line))
            except json.JSONDecodeError:
                print(f"[ログ警告] 壊れた行を無視: {os.path.basename(path)}")
    return rows


def _read_metadata(folder):
    meta = {}
    path = os.path.join(folder, "metadata.txt")
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                if "=" in line:
                    k, v = line.strip().split("=", 1)
                    meta[k] = v
    return meta


def _image_time_from_name(fname, start_time_str):
    """Elapsed seconds from a snapshot_HHMMSS_mmm.jpg name (pre-jsonl folders)."""
    if not start_time_str:
        return None
    try:
        hhmmss = os.path.splitext(fname)[0].split("_")[1]
        h, m, s = int(hhmmss[:2]), int(hhmmss[2:4]), int(hhmmss[4:6])
        st = datetime.datetime.strptime(start_time_str, "%Y-%m-%d %H:%M:%S")
        return (st.replace(hour=h, minute=m, second=s) - st).total_seconds()
    except Exception:
        return None


def collect_meeting_data(folder):
    """Gather everything a recording folder holds into one structure.

    Every renderer (markdown / HTML / DOCX) consumes this dict. None of them
    parses another renderer's output — that coupling would break the moment a
    heading or a table changes.
    """
    if not os.path.isdir(folder):
        raise FileNotFoundError(f"Folder not found: {folder}")

    audio_file = None
    for name in ("audio_main.mp3", "audio_main.wav"):
        path = os.path.join(folder, name)
        if os.path.exists(path):
            audio_file = path
            break
    if audio_file is None:
        raise RuntimeError(f"No audio file found in: {folder}")

    meta = _read_metadata(folder)

    # Per-role tracks, written only when both a mic and a speaker were captured.
    role_tracks = {}
    for role, fname in ((ROLE_SELF, ROLE_TRACK_SELF), (ROLE_OTHER, ROLE_TRACK_OTHER)):
        path = os.path.join(folder, fname)
        if os.path.exists(path) and os.path.getsize(path) > 0:
            role_tracks[role] = path

    # Image entries. snapshots.jsonl carries the exact timings, but the folder
    # is the source of truth for which images exist: any picture on disk that
    # the log missed still belongs in the report, timed from its filename.
    img_entries = []
    logged = set()
    for entry in _read_jsonl(os.path.join(folder, "snapshots.jsonl")):
        img_entries.append({
            "file": entry["file"],
            "time": entry.get("elapsed", 0.0),
            "type": entry.get("type", "auto"),
        })
        logged.add(entry["file"])
    start_str = meta.get("START_TIME_STR")
    missing = 0
    for fname in sorted(f for f in os.listdir(folder)
                        if f.lower().endswith((".jpg", ".jpeg", ".png"))):
        if fname in logged:
            continue
        t = _image_time_from_name(fname, start_str)
        if t is not None:
            img_entries.append({"file": fname, "time": t,
                                "type": "manual" if fname.startswith("manual_")
                                        else "auto"})
            missing += 1
    if missing and logged:
        print(f"[画像] ログに無い画像 {missing}枚をファイル名から復元しました")
    img_entries.sort(key=lambda x: x["time"])

    # OCR text is cached so re-running post-processing never redoes it
    ocr_by_file = {e.get("file"): e.get("text", "")
                   for e in _read_jsonl(os.path.join(folder, "ocr.jsonl"))}
    for ent in img_entries:
        ent["ocr"] = ocr_by_file.get(ent["file"], "")

    lang = meta.get("LANGUAGE", "ja")
    return {
        "folder": folder,
        "meta": meta,
        "meeting_name": meta.get("MEETING_NAME", ""),
        "start_time_str": meta.get("START_TIME_STR", "Unknown"),
        "audio_source": meta.get("AUDIO_SOURCE", "-"),
        "language": lang,
        "auto_mode": lang == "auto",
        "audio_file": audio_file,
        "audio_name": os.path.basename(audio_file),
        "role_tracks": role_tracks,
        "lang_segments": _read_jsonl(os.path.join(folder, "language_segments.jsonl")),
        "images": img_entries,
        "markers": sorted(_read_jsonl(os.path.join(folder, "markers.jsonl")),
                          key=lambda m: m.get("elapsed", 0.0)),
        "duration": 0.0,
        "detected_langs": [],
        "lines": [],
        "summary": None,
    }


# ---------------------------------------------------------------- Transcribing

class _WhisperEngine:
    """One interface over faster-whisper and the moonshine-voice fallback."""

    def __init__(self, prompt, report):
        self.prompt = prompt
        self.model = None
        self.name = "moonshine-voice"
        try:
            from faster_whisper import WhisperModel
            cfg = load_app_settings()
            model_name = cfg.get("whisper_model", "large-v3-turbo")
            device = cfg.get("whisper_device", "cpu")
            compute = cfg.get("whisper_compute", "int8")
            report(None, f"モデル読み込み中: {model_name} ({device})")
            self.model = WhisperModel(model_name, device=device, compute_type=compute)
            self.name = f"faster-whisper ({model_name})"
        except ImportError:
            report(None, "faster-whisper 未導入 → moonshine-voice にフォールバック")

    def detect_language(self, clip):
        return detect_ja_en(self.model, clip) if self.model is not None else "ja"

    def transcribe(self, clip, lang, sr):
        """Yield (start_seconds_within_clip, text) for one audio window."""
        if self.model is not None:
            segments, _info = self.model.transcribe(
                clip, language=lang, initial_prompt=self.prompt,
                vad_filter=True, beam_size=5)
            for seg in segments:
                text = seg.text.strip()
                if text:
                    yield (seg.start or 0.0), text
            return
        from moonshine_voice import Transcriber, get_model_for_language
        model_path, arch = get_model_for_language(
            "ja" if lang in (None, "auto") else lang)
        transcriber = Transcriber(model_path=model_path, model_arch=arch)
        try:
            transcript = transcriber.transcribe_without_streaming(clip.tolist(), sr)
        finally:
            transcriber.close()
        for line in transcript.lines:
            if line.text.strip():
                yield (line.start_time or 0.0), line.text.strip()


def _lang_at(lang_segments, elapsed, default_lang):
    """Language in force at a given elapsed second, per the switch log."""
    lang = default_lang
    for seg in sorted(lang_segments, key=lambda s: s.get("elapsed", 0.0)):
        if seg.get("elapsed", 0.0) <= elapsed:
            lang = seg.get("lang", lang)
        else:
            break
    return lang


def transcribe_meeting(data, report, span=(0.10, 0.80)):
    """Transcribe every audio track of a meeting into data["lines"].

    When per-role tracks exist the mixed audio is skipped: transcribing it as
    well would cover the same speech a third time for no benefit. Each track is
    cut into VAD speech windows first, so silence — most of a role track, since
    only one side speaks at a time — costs nothing.
    """
    terms = load_glossary()
    prompt = build_whisper_prompt(terms)
    print(f"Glossary terms: {len(terms)}")

    if data["role_tracks"]:
        tracks = [(path, role) for role, path in sorted(data["role_tracks"].items())]
        report(None, f"話者別トラック: {' / '.join(r for _, r in tracks)}")
    else:
        tracks = [(data["audio_file"], None)]

    engine = _WhisperEngine(prompt, report)
    data["engine"] = engine.name

    lines = []
    detected = set()
    auto_mode = data["auto_mode"]
    default_lang = "ja" if auto_mode else data["language"]
    lo, hi = span
    per_track = (hi - lo) / len(tracks)

    for t_idx, (path, speaker) in enumerate(tracks):
        base = lo + per_track * t_idx
        tag = f"[{speaker}] " if speaker else ""
        report(base, f"{tag}音声をデコード中...")
        audio, sr = decode_audio_16k_mono(path)
        duration = len(audio) / sr
        data["duration"] = max(data["duration"], duration)

        windows = [w for w in iter_speech_windows(audio, sr)
                   if (w[1] - w[0]) >= sr * 0.5]
        spoken = sum((e - s) / sr for s, e in windows) or 1.0
        report(base, f"{tag}発話区間 {len(windows)}件 / {spoken:.0f}秒")

        done = 0.0
        for start_sample, end_sample in windows:
            clip = audio[start_sample:end_sample]
            start_s = start_sample / sr
            lang = (engine.detect_language(clip) if auto_mode
                    else _lang_at(data["lang_segments"], start_s, default_lang))
            detected.add(lang)
            report(base + per_track * (done / spoken),
                   f"{tag}文字起こし {done:.0f}/{spoken:.0f}秒 [{start_s:.0f}s〜] {lang}")
            for offset, text in engine.transcribe(clip, lang, sr):
                lines.append({"start": start_s + offset, "text": text,
                              "speaker": speaker})
                # faster-whisper decodes lazily as segments are consumed, so a
                # cancel lands mid-window instead of waiting 25 seconds for one.
                report.check()
            done += (end_sample - start_sample) / sr
        del audio

    corrected = apply_glossary(lines, terms)
    if corrected:
        print(f"用語辞書による補正: {corrected}行")

    # Stable order: same instant sorts 相手 before 自分 so overlaps read the same
    # way every run.
    rank = {ROLE_OTHER: 0, ROLE_SELF: 1}
    lines.sort(key=lambda l: (l["start"], rank.get(l.get("speaker"), 9)))
    data["lines"] = lines
    data["detected_langs"] = sorted(detected)
    report(hi, f"文字起こし完了: {len(lines)}行")
    return lines


# ------------------------------------------------------------------ Rendering

def fmt_timestamp(seconds):
    """mm:ss, widening to h:mm:ss once a meeting passes the hour."""
    total = int(max(0.0, seconds))
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def build_timeline(data):
    """Merge transcript lines, screenshots and markers into one ordered stream."""
    events = []
    for line in data["lines"]:
        events.append({"t": line["start"], "kind": "line", "data": line})
    for img in data["images"]:
        events.append({"t": img["time"], "kind": "image", "data": img})
    for marker in data["markers"]:
        events.append({"t": marker.get("elapsed", 0.0), "kind": "marker",
                       "data": marker})
    # At an identical timestamp: marker, then the slide it refers to, then speech
    order = {"marker": 0, "image": 1, "line": 2}
    events.sort(key=lambda e: (e["t"], order.get(e["kind"], 9)))
    return events


def meeting_facts(data):
    """The metadata table shared by every renderer, as (label, value) pairs."""
    facts = []
    if data["meeting_name"]:
        facts.append(("会議名", data["meeting_name"]))
    facts.append(("開始時刻", data["start_time_str"]))
    dur = data["duration"]
    facts.append(("録音時間", f"{dur:.0f}秒 ({dur / 60:.1f}分)"))
    facts.append(("画像数", str(len(data["images"]))))
    if data["markers"]:
        facts.append(("マーカー", f"{len(data['markers'])}件"))
    facts.append(("音声ソース", data["audio_source"]))
    if data["role_tracks"]:
        facts.append(("話者分離", " / ".join(sorted(data["role_tracks"]))))
    if data["auto_mode"]:
        detected = "/".join(data["detected_langs"]) if data["detected_langs"] else "-"
        facts.append(("言語", f"自動 (ja/en) — 検出: {detected}"))
    else:
        facts.append(("言語", data["language"]))
    if data["lang_segments"] and not data["auto_mode"]:
        facts.append(("言語切替", ", ".join(
            f"{s.get('label', s.get('lang'))} ({s.get('elapsed', 0):.0f}s〜)"
            for s in data["lang_segments"])))
    if data.get("engine"):
        facts.append(("文字起こし", data["engine"]))
    return facts


def render_markdown(data, path):
    """Write meeting_report.md from collected data."""
    with open(path, "w", encoding="utf-8") as f:
        title = data["meeting_name"] or data["start_time_str"]
        f.write(f"# 会議記録 - {title}\n\n")
        f.write("| 項目 | 値 |\n|---|---|\n")
        for label, value in meeting_facts(data):
            f.write(f"| {label} | {value} |\n")
        f.write("\n")

        summary = data.get("summary")
        if summary:
            f.write(render_summary_markdown(summary))

        f.write("---\n\n## 記録\n\n")
        for event in build_timeline(data):
            ts = fmt_timestamp(event["t"])
            item = event["data"]
            if event["kind"] == "image":
                tag = "自動" if item.get("type") == "auto" else "手動"
                f.write(f"### [{ts}] 画面キャプチャ ({tag})\n\n")
                f.write(f"![{item['file']}]({item['file']})\n\n")
                if item.get("ocr"):
                    f.write("<details><summary>スライド内テキスト (OCR)</summary>\n\n")
                    f.write("```\n" + item["ocr"].strip() + "\n```\n\n")
                    f.write("</details>\n\n")
            elif event["kind"] == "marker":
                label = item.get("label") or "重要"
                f.write(f"> ⭐ **[{ts}] {label}**\n\n")
            else:
                speaker = item.get("speaker")
                who = f" {speaker}:" if speaker else ""
                f.write(f"**[{ts}]{who}** {item['text']}\n\n")

        f.write("---\n\n")
        f.write("*Generated by GijirokuStudio*\n")
    return path


def render_summary_markdown(summary):
    """The AI summary block, shared by the markdown and DOCX renderers."""
    out = ["## 📋 サマリ\n\n"]
    if summary.get("summary"):
        out.append(summary["summary"].strip() + "\n\n")
    if summary.get("decisions"):
        out.append("### 決定事項\n\n")
        out.extend(f"- {d}\n" for d in summary["decisions"])
        out.append("\n")
    if summary.get("action_items"):
        out.append("### アクションアイテム\n\n")
        out.append("| タスク | 担当 | 期限 |\n|---|---|---|\n")
        for item in summary["action_items"]:
            out.append(f"| {item.get('task', '')} | {item.get('owner') or '-'} "
                       f"| {item.get('due') or '-'} |\n")
        out.append("\n")
    if summary.get("open_issues"):
        out.append("### 未決事項\n\n")
        out.extend(f"- {q}\n" for q in summary["open_issues"])
        out.append("\n")
    if summary.get("raw"):
        out.append("### AI要約（未整形）\n\n```\n" + summary["raw"].strip() + "\n```\n\n")
    return "".join(out)


# ----------------------------------------------------------------- Slide OCR

# Windows ships a Japanese OCR engine (Windows.Media.Ocr). Reaching it needs
# WinRT, and no WinRT binding installs on Python 3.13 — winsdk, which winocr
# depends on, has no 3.13 wheel — so we drive it through PowerShell instead.
# One process handles every image so the WinRT setup cost is paid once.
OCR_PS_SCRIPT = r"""
param([Parameter(Mandatory=$true)][string]$ListPath,
      [Parameter(Mandatory=$true)][string]$OutPath,
      [string]$Lang = 'ja')

$ErrorActionPreference = 'Stop'
Add-Type -AssemblyName System.Runtime.WindowsRuntime | Out-Null
$null = [Windows.Media.Ocr.OcrEngine, Windows.Foundation, ContentType=WindowsRuntime]
$null = [Windows.Graphics.Imaging.BitmapDecoder, Windows.Foundation, ContentType=WindowsRuntime]
$null = [Windows.Storage.StorageFile, Windows.Foundation, ContentType=WindowsRuntime]
$null = [Windows.Globalization.Language, Windows.Foundation, ContentType=WindowsRuntime]

$asTaskGeneric = ([System.WindowsRuntimeSystemExtensions].GetMethods() | Where-Object {
    $_.Name -eq 'AsTask' -and $_.GetParameters().Count -eq 1 -and
    $_.GetParameters()[0].ParameterType.Name -eq 'IAsyncOperation`1' })[0]

function Await($task, $type) {
    $m = $asTaskGeneric.MakeGenericMethod($type)
    $t = $m.Invoke($null, @($task))
    $t.Wait(-1) | Out-Null
    $t.Result
}

$engine = $null
try {
    $engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromLanguage(
        [Windows.Globalization.Language]::new($Lang))
} catch {}
if ($null -eq $engine) {
    $engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromUserProfileLanguages()
}
if ($null -eq $engine) { Write-Error 'NO_OCR_ENGINE'; exit 2 }

$writer = New-Object System.IO.StreamWriter($OutPath, $false, (New-Object System.Text.UTF8Encoding($false)))
foreach ($path in [System.IO.File]::ReadAllLines($ListPath, [System.Text.Encoding]::UTF8)) {
    if ([string]::IsNullOrWhiteSpace($path)) { continue }
    $text = ''
    $err = ''
    try {
        $file = Await ([Windows.Storage.StorageFile]::GetFileFromPathAsync($path)) ([Windows.Storage.StorageFile])
        $stream = Await ($file.OpenAsync([Windows.Storage.FileAccessMode]::Read)) ([Windows.Storage.Streams.IRandomAccessStream])
        $decoder = Await ([Windows.Graphics.Imaging.BitmapDecoder]::CreateAsync($stream)) ([Windows.Graphics.Imaging.BitmapDecoder])
        $bitmap = Await ($decoder.GetSoftwareBitmapAsync()) ([Windows.Graphics.Imaging.SoftwareBitmap])
        $result = Await ($engine.RecognizeAsync($bitmap)) ([Windows.Media.Ocr.OcrResult])
        $sb = New-Object System.Text.StringBuilder
        foreach ($line in $result.Lines) { [void]$sb.AppendLine($line.Text) }
        $text = $sb.ToString()
        $stream.Dispose()
    } catch {
        # Surfaced to the caller — a silently empty result looks like a blank
        # slide, which is a very different problem from an unreadable file.
        $err = $_.Exception.Message
    }
    $writer.WriteLine((@{ file = $path; text = $text; error = $err } | ConvertTo-Json -Compress))
}
$writer.Close()
"""

# The Windows engine emits one "word" per glyph for Japanese, so its output
# arrives as "売 上 高 : 1 , 240". Those spaces are segmentation artifacts, not
# slide content — but Latin words genuinely are space-separated, so the cleanup
# only closes gaps where at least one side is Japanese, plus digit-group and
# punctuation artifacts.
_CJK = r"぀-ヿ㐀-䶿一-鿿＀-￯"
_SP = r"[ 　]"
_OCR_FIXES = [
    (re.compile(f"(?<=[{_CJK}]){_SP}+(?=[{_CJK}])"), ""),          # 売 上 -> 売上
    (re.compile(f"(?<=[{_CJK}]){_SP}+(?=[0-9A-Za-z(（\\[])"), ""),  # 第 3 -> 第3
    (re.compile(f"(?<=[0-9A-Za-z)）\\]%％]){_SP}+(?=[{_CJK}])"), ""),  # 42 社 -> 42社
    (re.compile(rf"(?<=\d){_SP}*([,.]){_SP}*(?=\d)"), r"\1"),      # 1 , 240 -> 1,240
    (re.compile(f"{_SP}+(?=[,.:;%％)）\\]])"), ""),                  # 118 % -> 118%
    (re.compile(f"(?<=[(（\\[]){_SP}+"), ""),                       # ( 前年比 -> (前年比
    (re.compile(f"(?<=[:：]){_SP}+(?=[0-9{_CJK}])"), ""),           # 高: 1 -> 高:1
]


def tidy_ocr_text(text):
    """Clean OCR output: close glyph-level gaps and drop empty lines."""
    lines = []
    for line in (text or "").splitlines():
        for pattern, repl in _OCR_FIXES:
            line = pattern.sub(repl, line)
        line = line.strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def _ocr_windows(paths, lang):
    """Run the Windows OCR engine over a batch of images. {path: text}."""
    if os.name != "nt":
        raise RuntimeError("Windows OCR は Windows 専用です")
    # StorageFile.GetFileFromPathAsync rejects forward slashes, and tkinter's
    # askdirectory() hands back exactly that on Windows.
    paths = [os.path.normpath(os.path.abspath(p)) for p in paths]
    tmpdir = tempfile.mkdtemp(prefix="gjs_ocr_")
    try:
        script = os.path.join(tmpdir, "ocr.ps1")
        listing = os.path.join(tmpdir, "images.txt")
        result = os.path.join(tmpdir, "out.jsonl")
        with open(script, "w", encoding="utf-8-sig") as f:
            f.write(OCR_PS_SCRIPT)
        with open(listing, "w", encoding="utf-8") as f:
            f.write("\n".join(paths))
        proc = subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive",
             "-ExecutionPolicy", "Bypass", "-File", script,
             "-ListPath", listing, "-OutPath", result, "-Lang", lang],
            capture_output=True, text=True, timeout=600,
            creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0)
        if not os.path.exists(result):
            raise RuntimeError((proc.stderr or "PowerShell 実行失敗").strip()[:200])
        out = {}
        with open(result, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                row = json.loads(line)
                if row.get("error"):
                    print(f"[OCR警告] {os.path.basename(row['file'])}: {row['error']}")
                out[row["file"]] = row.get("text", "")
        return out
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def _ocr_rapidocr(paths, _lang):
    """Fallback engine. Its default model is Chinese, so kanji can come back as
    simplified variants — usable for numbers and Latin text, weaker on Japanese."""
    from rapidocr_onnxruntime import RapidOCR
    engine = RapidOCR()
    out = {}
    for path in paths:
        try:
            result, _elapsed = engine(path)
            out[path] = "\n".join(item[1] for item in (result or []))
        except Exception as e:
            print(f"[OCR警告] {os.path.basename(path)}: {e}")
            out[path] = ""
    return out


OCR_BACKENDS = {"windows": _ocr_windows, "rapidocr": _ocr_rapidocr}


def ocr_images(data, report):
    """Extract slide text into data["images"][*]["ocr"], caching to ocr.jsonl.

    Runs in post-processing, never during recording — the CPU belongs to the
    meeting. Results are cached so re-running post-processing never redoes it.
    """
    cfg = load_app_settings()
    if not cfg.get("ocr_enabled", False):
        return 0
    folder = data["folder"]
    done = {e.get("file") for e in _read_jsonl(os.path.join(folder, "ocr.jsonl"))}
    todo = [e for e in data["images"] if e["file"] not in done]
    if not todo:
        return 0

    paths = [os.path.join(folder, e["file"]) for e in todo]
    paths = [p for p in paths if os.path.exists(p)]
    if not paths:
        return 0

    wanted = str(cfg.get("ocr_backend", "auto")).lower()
    order = ([wanted] if wanted in OCR_BACKENDS else ["windows", "rapidocr"])
    lang = str(cfg.get("ocr_lang", "ja"))

    results = None
    for name in order:
        try:
            report(None, f"スライドOCR ({name}): {len(paths)}枚")
            results = OCR_BACKENDS[name](paths, lang)
            break
        except Exception as e:
            report(None, f"[OCR] {name} 利用不可: {e}")
    if not results:
        report(None, "OCR をスキップしました")
        return 0

    by_name = {e["file"]: e for e in data["images"]}
    hits = 0
    with open(os.path.join(folder, "ocr.jsonl"), "a", encoding="utf-8") as f:
        for path, raw in results.items():
            fname = os.path.basename(path)
            text = tidy_ocr_text(raw)
            if fname in by_name:
                by_name[fname]["ocr"] = text
            f.write(json.dumps({"file": fname, "text": text},
                               ensure_ascii=False) + "\n")
            if text:
                hits += 1
    report(None, f"OCR 完了: {hits}/{len(paths)}枚から文字を抽出")
    return hits


# ------------------------------------------------------- HTML / DOCX export

HTML_STYLE = """
:root { color-scheme: light dark; --fg:#1f2937; --bg:#ffffff; --mut:#6b7280;
        --line:#e5e7eb; --accent:#0ea5e9; --mark:#f59e0b; --card:#f9fafb; }
@media (prefers-color-scheme: dark) {
  :root { --fg:#e5e7eb; --bg:#111827; --mut:#9ca3af; --line:#374151;
          --accent:#38bdf8; --mark:#fbbf24; --card:#1f2937; } }
* { box-sizing: border-box; }
body { margin:0; padding:0 0 4rem; background:var(--bg); color:var(--fg);
       font-family:"BIZ UDPゴシック","Yu Gothic UI","Hiragino Sans",sans-serif;
       line-height:1.75; }
.wrap { max-width:52rem; margin:0 auto; padding:0 1.2rem; }
h1 { font-size:1.6rem; margin:1.5rem 0 .5rem; }
h2 { font-size:1.2rem; margin:2rem 0 .6rem; border-bottom:2px solid var(--line);
     padding-bottom:.3rem; }
h3 { font-size:1rem; margin:1.4rem 0 .4rem; color:var(--mut); }
table { border-collapse:collapse; width:100%; margin:.6rem 0; font-size:.9rem; }
th,td { border:1px solid var(--line); padding:.4rem .6rem; text-align:left;
        vertical-align:top; }
th { background:var(--card); white-space:nowrap; }
.meta td:first-child { white-space:nowrap; color:var(--mut); width:8rem; }
.scroll { overflow-x:auto; }
.player { position:sticky; top:0; z-index:10; background:var(--bg);
          border-bottom:1px solid var(--line); padding:.6rem 0; }
.player audio { width:100%; }
.tools { display:flex; gap:.5rem; align-items:center; margin:.5rem 0 0; }
.tools input { flex:1; padding:.4rem .6rem; border:1px solid var(--line);
               border-radius:.3rem; background:var(--bg); color:var(--fg);
               font-size:.9rem; }
.line { display:flex; gap:.6rem; padding:.15rem 0; scroll-margin-top:6rem; }
.line.hide { display:none; }
.t { flex:none; font-variant-numeric:tabular-nums; font-size:.8rem;
     color:var(--accent); background:none; border:0; padding:.1rem .2rem;
     cursor:pointer; font-family:inherit; }
.t:hover { text-decoration:underline; }
.who { flex:none; font-size:.8rem; color:var(--mut); min-width:2.6rem; }
.said { flex:1; }
.marker { margin:.8rem 0; padding:.4rem .8rem; border-left:4px solid var(--mark);
          background:var(--card); font-weight:bold; }
figure { margin:1.2rem 0; }
figure img { max-width:100%; border:1px solid var(--line); border-radius:.3rem;
             display:block; }
figcaption { font-size:.8rem; color:var(--mut); margin-top:.3rem; }
details { margin:.4rem 0; font-size:.85rem; }
details pre { white-space:pre-wrap; background:var(--card); padding:.6rem;
              border-radius:.3rem; margin:.4rem 0 0; }
.hit { background:rgba(245,158,11,.25); }
footer { color:var(--mut); font-size:.8rem; margin-top:2.5rem;
         border-top:1px solid var(--line); padding-top:.8rem; }
@media print { .player, .tools { position:static; } .t { color:var(--mut); } }
"""

HTML_SCRIPT = """
(function () {
  var audio = document.getElementById('player');
  document.querySelectorAll('[data-t]').forEach(function (el) {
    el.addEventListener('click', function () {
      if (!audio) return;
      audio.currentTime = parseFloat(el.dataset.t) || 0;
      audio.play();
    });
  });
  var box = document.getElementById('q');
  if (box) {
    box.addEventListener('input', function () {
      var q = box.value.trim();
      document.querySelectorAll('.line').forEach(function (el) {
        var said = el.querySelector('.said');
        var text = said.textContent;
        el.classList.toggle('hide', q !== '' && text.indexOf(q) === -1);
        said.innerHTML = '';
        if (q === '' || text.indexOf(q) === -1) { said.textContent = text; return; }
        var i = 0, at;
        while ((at = text.indexOf(q, i)) !== -1) {
          said.appendChild(document.createTextNode(text.slice(i, at)));
          var m = document.createElement('mark');
          m.className = 'hit';
          m.textContent = q;
          said.appendChild(m);
          i = at + q.length;
        }
        said.appendChild(document.createTextNode(text.slice(i)));
      });
    });
  }
})();
"""


def _data_uri(path, mime):
    """base64 data: URI for a file, or None if it can't be read."""
    try:
        with open(path, "rb") as f:
            return f"data:{mime};base64," + base64.b64encode(f.read()).decode("ascii")
    except OSError as e:
        print(f"[HTML警告] 埋め込めません {os.path.basename(path)}: {e}")
        return None


def _shrink_audio_for_embed(src):
    """Re-encode to 64kbps mono for embedding — a 192kbps hour is ~115MB in base64."""
    tmp = f"{src}._embed.mp3"
    flags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
    subprocess.run([FFMPEG_PATH, "-y", "-i", src, "-ac", "1", "-b:a", "64k", tmp],
                   stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                   check=True, creationflags=flags)
    return tmp


def render_html(data, path, embed_audio=False, embed_images=True):
    """Write a shareable single-page report from collected data.

    Images are inlined so the file survives being emailed on its own. Audio is
    referenced by relative path unless `embed_audio` is set: an hour of 192kbps
    MP3 becomes ~115MB once base64-encoded, which no browser enjoys opening.
    """
    esc = html.escape
    folder = data["folder"]
    title = data["meeting_name"] or data["start_time_str"]
    out = []
    add = out.append

    add("<!DOCTYPE html>\n<html lang=\"ja\">\n<head>\n<meta charset=\"utf-8\">\n")
    add('<meta name="viewport" content="width=device-width,initial-scale=1">\n')
    add(f"<title>会議記録 - {esc(title)}</title>\n<style>{HTML_STYLE}</style>\n")
    add("</head>\n<body>\n<div class=\"wrap\">\n")
    add(f"<h1>会議記録 - {esc(title)}</h1>\n")

    audio_src = data["audio_name"]
    tmp_audio = None
    if embed_audio:
        try:
            tmp_audio = _shrink_audio_for_embed(data["audio_file"])
            uri = _data_uri(tmp_audio, "audio/mpeg")
            if uri:
                audio_src = uri
        except Exception as e:
            print(f"[HTML警告] 音声を埋め込めません: {e}")
    add('<div class="player">\n')
    add(f'<audio id="player" controls preload="none" src="{audio_src}"></audio>\n')
    add('<div class="tools"><input id="q" type="search" '
        'placeholder="発言を検索（クリックでその時刻から再生）"></div>\n')
    add("</div>\n")
    if tmp_audio:
        try:
            os.remove(tmp_audio)
        except OSError:
            pass

    add('<div class="scroll"><table class="meta"><tbody>\n')
    for label, value in meeting_facts(data):
        add(f"<tr><td>{esc(label)}</td><td>{esc(str(value))}</td></tr>\n")
    add("</tbody></table></div>\n")

    summary = data.get("summary")
    if summary:
        add('<h2>📋 サマリ</h2>\n')
        if summary.get("summary"):
            add(f"<p>{esc(summary['summary']).replace(chr(10), '<br>')}</p>\n")
        if summary.get("decisions"):
            add("<h3>決定事項</h3>\n<ul>\n")
            add("".join(f"<li>{esc(d)}</li>\n" for d in summary["decisions"]))
            add("</ul>\n")
        if summary.get("action_items"):
            add("<h3>アクションアイテム</h3>\n<div class=\"scroll\"><table>\n")
            add("<thead><tr><th>タスク</th><th>担当</th><th>期限</th></tr></thead><tbody>\n")
            for item in summary["action_items"]:
                add(f"<tr><td>{esc(item.get('task', ''))}</td>"
                    f"<td>{esc(item.get('owner') or '-')}</td>"
                    f"<td>{esc(item.get('due') or '-')}</td></tr>\n")
            add("</tbody></table></div>\n")
        if summary.get("open_issues"):
            add("<h3>未決事項</h3>\n<ul>\n")
            add("".join(f"<li>{esc(q)}</li>\n" for q in summary["open_issues"]))
            add("</ul>\n")
        if summary.get("raw"):
            add("<h3>AI要約（未整形）</h3>\n"
                f"<pre>{esc(summary['raw'])}</pre>\n")

    add("<h2>記録</h2>\n")
    for event in build_timeline(data):
        ts = fmt_timestamp(event["t"])
        item = event["data"]
        seek = f'{event["t"]:.1f}'
        if event["kind"] == "image":
            src = item["file"]
            if embed_images:
                uri = _data_uri(os.path.join(folder, item["file"]), "image/jpeg")
                if uri:
                    src = uri
            tag = "自動" if item.get("type") == "auto" else "手動"
            add("<figure>\n")
            add(f'<button class="t" data-t="{seek}">[{ts}]</button> '
                f'<span class="who">画面キャプチャ ({tag})</span>\n')
            add(f'<img src="{src}" alt="{esc(item["file"])}" loading="lazy">\n')
            add(f"<figcaption>{esc(item['file'])}</figcaption>\n")
            if item.get("ocr"):
                add("<details><summary>スライド内テキスト (OCR)</summary>"
                    f"<pre>{esc(item['ocr'].strip())}</pre></details>\n")
            add("</figure>\n")
        elif event["kind"] == "marker":
            label = item.get("label") or "重要"
            add(f'<div class="marker"><button class="t" data-t="{seek}">[{ts}]</button> '
                f'⭐ {esc(label)}</div>\n')
        else:
            who = item.get("speaker") or ""
            add('<div class="line">'
                f'<button class="t" data-t="{seek}">[{ts}]</button>'
                f'<span class="who">{esc(who)}</span>'
                f'<span class="said">{esc(item["text"])}</span></div>\n')

    add("<footer>Generated by GijirokuStudio</footer>\n</div>\n")
    add(f"<script>{HTML_SCRIPT}</script>\n</body>\n</html>\n")

    with open(path, "w", encoding="utf-8") as f:
        f.write("".join(out))
    size_mb = os.path.getsize(path) / (1024 * 1024)
    if size_mb > 40:
        print(f"[HTML警告] ファイルが大きすぎます ({size_mb:.0f}MB)。"
              "画像埋め込みを無効にするか枚数を減らしてください")
    return path, size_mb


def render_docx(data, path):
    """Write a .docx report. Requires python-docx; raises RuntimeError without it."""
    try:
        import docx
        from docx.shared import Pt, Inches
    except ImportError:
        raise RuntimeError("python-docx が未導入です（pipenv install python-docx）")

    folder = data["folder"]
    doc = docx.Document()
    doc.add_heading(f"会議記録 - {data['meeting_name'] or data['start_time_str']}", 0)

    facts = meeting_facts(data)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in facts:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    summary = data.get("summary")
    if summary:
        doc.add_heading("サマリ", level=1)
        if summary.get("summary"):
            doc.add_paragraph(summary["summary"])
        if summary.get("decisions"):
            doc.add_heading("決定事項", level=2)
            for d in summary["decisions"]:
                doc.add_paragraph(d, style="List Bullet")
        if summary.get("action_items"):
            doc.add_heading("アクションアイテム", level=2)
            at = doc.add_table(rows=1, cols=3)
            at.style = "Light Grid Accent 1"
            for i, head in enumerate(("タスク", "担当", "期限")):
                at.rows[0].cells[i].text = head
            for item in summary["action_items"]:
                cells = at.add_row().cells
                cells[0].text = item.get("task", "")
                cells[1].text = item.get("owner") or "-"
                cells[2].text = item.get("due") or "-"
        if summary.get("open_issues"):
            doc.add_heading("未決事項", level=2)
            for q in summary["open_issues"]:
                doc.add_paragraph(q, style="List Bullet")

    doc.add_heading("記録", level=1)
    for event in build_timeline(data):
        ts = fmt_timestamp(event["t"])
        item = event["data"]
        if event["kind"] == "image":
            tag = "自動" if item.get("type") == "auto" else "手動"
            doc.add_paragraph(f"[{ts}] 画面キャプチャ ({tag})")
            img_path = os.path.join(folder, item["file"])
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                except Exception as e:
                    print(f"[DOCX警告] 画像を挿入できません {item['file']}: {e}")
            if item.get("ocr"):
                para = doc.add_paragraph(item["ocr"].strip())
                para.runs[0].font.size = Pt(8)
        elif event["kind"] == "marker":
            para = doc.add_paragraph(f"⭐ [{ts}] {item.get('label') or '重要'}")
            para.runs[0].bold = True
        else:
            who = f" {item['speaker']}:" if item.get("speaker") else ""
            para = doc.add_paragraph()
            run = para.add_run(f"[{ts}]{who} ")
            run.bold = True
            para.add_run(item["text"])

    doc.save(path)
    return path


# ------------------------------------------------------------------ AI summary

# The contract every provider must satisfy. Both back-ends are told to emit
# exactly this, so the report renderer never has to guess at free-form prose.
SUMMARY_SCHEMA = {
    "type": "object",
    "properties": {
        "summary": {"type": "string"},
        "decisions": {"type": "array", "items": {"type": "string"}},
        "action_items": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "task": {"type": "string"},
                    "owner": {"type": ["string", "null"]},
                    "due": {"type": ["string", "null"]},
                },
                "required": ["task", "owner", "due"],
                "additionalProperties": False,
            },
        },
        "open_issues": {"type": "array", "items": {"type": "string"}},
    },
    "required": ["summary", "decisions", "action_items", "open_issues"],
    "additionalProperties": False,
}

SUMMARY_SYSTEM = """あなたは会議の議事録作成者です。文字起こしから事実だけを抽出します。

- 推測で補わない。文字起こしにない内容は書かない
- 決定事項は「決まったこと」だけ。検討中の案は未決事項へ
- アクションアイテムは担当者と期限を文字起こしから読み取る。不明なら null
- 話者が「自分」「相手」と付いている場合、担当者の判断に使う
- ⭐ が付いた箇所は発言者が重要と判断した部分なので優先的に拾う
- 出力は日本語"""

SUMMARY_INSTRUCTION = """次の会議の文字起こしから、サマリ・決定事項・アクション\
アイテム・未決事項を抽出してください。"""


def transcript_for_summary(data, max_chars=None):
    """Flatten the transcript for an LLM, marking bookmarked moments with ⭐."""
    marks = [m.get("elapsed", 0.0) for m in data.get("markers", [])]
    out = []
    for line in data["lines"]:
        start = line["start"]
        star = "⭐ " if any(abs(start - m) <= 60.0 for m in marks) else ""
        speaker = f"{line['speaker']}: " if line.get("speaker") else ""
        out.append(f"[{fmt_timestamp(start)}] {star}{speaker}{line['text']}")
    text = "\n".join(out)
    if max_chars and len(text) > max_chars:
        text = text[:max_chars]
    return text


def _parse_summary_json(text):
    """Parse a model's reply into the summary dict, degrading rather than failing.

    Tries strict JSON, then a fenced ```json block, then gives up and returns the
    raw text so the report can still show something.
    """
    if not text:
        return None
    try:
        return _coerce_summary(json.loads(text))
    except (json.JSONDecodeError, TypeError):
        pass
    match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.S)
    if not match:
        match = re.search(r"(\{.*\})", text, re.S)
    if match:
        try:
            return _coerce_summary(json.loads(match.group(1)))
        except (json.JSONDecodeError, TypeError):
            pass
    print("[要約警告] JSON として読めなかったため原文を掲載します")
    return {"summary": "", "decisions": [], "action_items": [],
            "open_issues": [], "raw": text}


def _coerce_summary(obj):
    """Normalize a parsed object to the schema, tolerating loose model output."""
    if not isinstance(obj, dict):
        raise TypeError("summary must be an object")
    items = []
    for item in obj.get("action_items") or []:
        if isinstance(item, dict):
            items.append({"task": str(item.get("task", "")).strip(),
                          "owner": item.get("owner") or None,
                          "due": item.get("due") or None})
        elif isinstance(item, str):
            items.append({"task": item, "owner": None, "due": None})
    as_list = lambda v: [str(x) for x in v] if isinstance(v, list) else []
    return {
        "summary": str(obj.get("summary") or "").strip(),
        "decisions": as_list(obj.get("decisions")),
        "action_items": [i for i in items if i["task"]],
        "open_issues": as_list(obj.get("open_issues")),
    }


def _merge_summaries(parts):
    """Concatenate map-stage results, de-duplicating while preserving order."""
    merged = {"summary": "", "decisions": [], "action_items": [], "open_issues": []}
    seen = {"decisions": set(), "open_issues": set(), "action_items": set()}
    texts = []
    for part in parts:
        if not part:
            continue
        if part.get("summary"):
            texts.append(part["summary"])
        for key in ("decisions", "open_issues"):
            for value in part.get(key, []):
                if value not in seen[key]:
                    seen[key].add(value)
                    merged[key].append(value)
        for item in part.get("action_items", []):
            key = (item["task"], item.get("owner"))
            if key not in seen["action_items"]:
                seen["action_items"].add(key)
                merged["action_items"].append(item)
    merged["summary"] = "\n\n".join(texts)
    return merged


def _chunk_transcript(text, max_chars):
    """Split on line boundaries so an utterance is never cut in half."""
    chunks, current = [], []
    size = 0
    for line in text.splitlines():
        if size + len(line) + 1 > max_chars and current:
            chunks.append("\n".join(current))
            current, size = [], 0
        current.append(line)
        size += len(line) + 1
    if current:
        chunks.append("\n".join(current))
    return chunks


def _summarize_claude(transcript, cfg, report):
    """One request to the Claude API — a meeting fits in the context window."""
    try:
        import anthropic
    except ImportError:
        raise RuntimeError("anthropic パッケージが未導入です（pipenv install anthropic）")

    model = cfg.get("summary_model") or "claude-opus-5"
    api_key = cfg.get("anthropic_api_key") or None
    client = anthropic.Anthropic(api_key=api_key) if api_key else anthropic.Anthropic()
    report(None, f"要約を生成中（Claude API / {model}）— 文字起こしを外部送信します")

    params = dict(
        model=model,
        max_tokens=32000,
        system=SUMMARY_SYSTEM,
        output_config={"format": {"type": "json_schema", "schema": SUMMARY_SCHEMA}},
        messages=[{"role": "user",
                   "content": f"{SUMMARY_INSTRUCTION}\n\n<transcript>\n{transcript}\n</transcript>"}],
    )

    def _run(with_fallback):
        # Claude Opus 5's safety classifiers can decline a request; the
        # server-side fallback re-runs it on another model in the same call.
        if with_fallback:
            with client.beta.messages.stream(
                    betas=["server-side-fallback-2026-07-01"],
                    fallbacks="default", **params) as stream:
                return stream.get_final_message()
        with client.messages.stream(**params) as stream:
            return stream.get_final_message()

    try:
        message = _run(True)
    except anthropic.BadRequestError as e:
        # Older account/SDK without the fallback beta — proceed without it.
        print(f"[要約] フォールバック無しで再試行: {e}")
        message = _run(False)

    if message.stop_reason == "refusal":
        raise RuntimeError("安全性判定により要約が拒否されました")
    text = next((b.text for b in message.content if b.type == "text"), "")
    return _parse_summary_json(text)


def _ollama_chat(url, model, system, user, report):
    """One Ollama /api/chat call with JSON mode forced."""
    payload = json.dumps({
        "model": model,
        "stream": False,
        "format": "json",
        "messages": [{"role": "system", "content": system},
                     {"role": "user", "content": user}],
    }).encode("utf-8")
    request = urllib.request.Request(
        url.rstrip("/") + "/api/chat", data=payload,
        headers={"Content-Type": "application/json"})
    # No default timeout on urllib, and a CPU-hosted model takes minutes.
    with urllib.request.urlopen(request, timeout=900) as response:
        body = json.loads(response.read().decode("utf-8"))
    return body.get("message", {}).get("content", "")


def _summarize_ollama(transcript, cfg, report):
    """Map-reduce over the transcript — local models have small context windows."""
    url = cfg.get("ollama_url") or "http://localhost:11434"
    model = cfg.get("summary_model") or "qwen3"
    chunks = _chunk_transcript(transcript, int(cfg.get("summary_chunk_chars", 6000)))
    report(None, f"要約を生成中（Ollama / {model}） 分割 {len(chunks)}件")

    schema_hint = ('必ず次の JSON だけを返してください: '
                   '{"summary": "...", "decisions": ["..."], '
                   '"action_items": [{"task": "...", "owner": null, "due": null}], '
                   '"open_issues": ["..."]}')
    parts = []
    for i, chunk in enumerate(chunks, 1):
        report(None, f"要約 {i}/{len(chunks)}")
        reply = _ollama_chat(url, model, SUMMARY_SYSTEM,
                             f"{SUMMARY_INSTRUCTION}\n{schema_hint}\n\n"
                             f"（会議の一部 {i}/{len(chunks)}）\n{chunk}", report)
        parts.append(_parse_summary_json(reply))

    merged = _merge_summaries(parts)
    if len(chunks) == 1:
        return merged
    # Reduce: fold the per-chunk extracts into one coherent set.
    report(None, "要約を統合中")
    reply = _ollama_chat(url, model, SUMMARY_SYSTEM,
                         f"次は同じ会議を分割して抽出した結果です。重複を除いて"
                         f"1つにまとめてください。\n{schema_hint}\n\n"
                         f"{json.dumps(merged, ensure_ascii=False, indent=1)}", report)
    return _parse_summary_json(reply) or merged


def summarize_meeting(data, report):
    """Attach an AI summary to the meeting data, or leave it None.

    A summarization failure never fails post-processing — the transcript and
    report are the deliverable; the summary is an addition to them.
    """
    cfg = load_app_settings()
    provider = str(cfg.get("summary_provider", "none")).lower()
    if provider in ("", "none", "off"):
        return None
    if not data["lines"]:
        report(None, "発言がないため要約をスキップ")
        return None

    transcript = transcript_for_summary(data)
    try:
        if provider == "claude":
            summary = _summarize_claude(transcript, cfg, report)
        elif provider == "ollama":
            summary = _summarize_ollama(transcript, cfg, report)
        else:
            report(None, f"[要約] 未知のプロバイダ: {provider}")
            return None
    except PostProcessCancelled:
        raise
    except Exception as e:
        report(None, f"[要約エラー] {e}（議事録の生成は続行します）")
        traceback.print_exc()
        return None

    if summary:
        summary["provider"] = provider
        data["summary"] = summary
        report(None, f"要約完了: 決定事項{len(summary.get('decisions', []))}件 "
                     f"/ アクション{len(summary.get('action_items', []))}件")
    return summary


# --------------------------------------------------------------- Entry point

def _extract_video_snapshots(video_path, folder, interval=5.0, threshold=10,
                             max_edge=1600, jpeg_quality=85,
                             progress=None, cancel=None):
    """Sample a video and keep frames whose dHash changed significantly."""
    frame_dir = os.path.join(folder, ".video_frames")
    os.makedirs(frame_dir)
    pattern = os.path.join(frame_dir, "frame_%06d.jpg")
    flags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
    args = [
        FFMPEG_PATH, "-y", "-loglevel", "error", "-i", video_path,
        "-vf", f"fps=1/{max(0.1, float(interval))}", "-q:v", "3", pattern,
    ]
    proc = subprocess.Popen(
        args, stdout=subprocess.DEVNULL, stderr=subprocess.PIPE,
        creationflags=flags)
    while proc.poll() is None:
        if cancel is not None and cancel():
            proc.terminate()
            try:
                proc.wait(timeout=5)
            except subprocess.TimeoutExpired:
                proc.kill()
            raise PostProcessCancelled()
        time.sleep(0.2)
    stderr = proc.stderr.read().decode("utf-8", errors="replace").strip()
    if proc.returncode != 0:
        detail = stderr.splitlines()[-1] if stderr else "フレーム抽出に失敗しました"
        raise RuntimeError(f"動画の画像を抽出できませんでした: {detail}")

    frames = sorted(os.path.join(frame_dir, name) for name in os.listdir(frame_dir)
                    if name.lower().endswith(".jpg"))
    kept = []
    last_hash = None
    try:
        for index, source in enumerate(frames):
            if cancel is not None and cancel():
                raise PostProcessCancelled()
            with Image.open(source) as opened:
                image = opened.convert("RGB")
                current_hash = imagehash.dhash(image)
                diff = 0 if last_hash is None else int(current_hash - last_hash)
                if last_hash is not None and diff <= threshold:
                    continue
                last_hash = current_hash
                if max_edge and max(image.size) > max_edge:
                    scale = max_edge / max(image.size)
                    image = image.resize(
                        (max(1, round(image.width * scale)),
                         max(1, round(image.height * scale))), Image.LANCZOS)
                elapsed = index * float(interval)
                filename = f"snapshot_video_{round(elapsed * 1000):09d}.jpg"
                image.save(os.path.join(folder, filename), "JPEG", quality=jpeg_quality)
            kept.append({
                "file": filename, "elapsed": round(elapsed, 3),
                "type": "video", "diff": diff,
            })
            if progress is not None and len(kept) % 10 == 0:
                progress(None, f"動画の画面変化を抽出中: {len(kept)}枚保存")
        if kept:
            with open(os.path.join(folder, "snapshots.jsonl"), "w", encoding="utf-8") as f:
                for entry in kept:
                    f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    finally:
        shutil.rmtree(frame_dir, ignore_errors=True)
    return len(kept)


def import_video_file(video_path, language="auto", progress=None, cancel=None,
                      capture_scenes=False, scene_interval=5.0,
                      scene_threshold=10, max_edge=1600, jpeg_quality=85):
    """Create a meeting folder from a video without modifying the source file.

    The video stream itself is not copied. ffmpeg extracts a standard
    ``audio_main.mp3`` so the existing transcription and report pipeline can be
    reused without special cases in renderers.
    """
    video_path = os.path.abspath(video_path)
    if not os.path.isfile(video_path):
        raise FileNotFoundError(f"動画ファイルが見つかりません: {video_path}")
    if not os.path.isfile(FFMPEG_PATH):
        raise FileNotFoundError(f"ffmpeg.exe が見つかりません: {FFMPEG_PATH}")
    if cancel is not None and cancel():
        raise PostProcessCancelled()

    def _progress(frac, message):
        print(message)
        if progress is not None:
            progress(frac, message)

    stem = os.path.splitext(os.path.basename(video_path))[0]
    safe = MeetingRecorderGUI._sanitize_name(stem) or "Video"
    now = datetime.datetime.now()
    base_name = f"{now.strftime('%Y%m%d_%H%M%S')}_{safe}"
    folder = os.path.join(BASE_DIR, base_name)
    suffix = 2
    while os.path.exists(folder):
        folder = os.path.join(BASE_DIR, f"{base_name}_{suffix}")
        suffix += 1
    os.makedirs(folder)

    audio_path = os.path.join(folder, "audio_main.mp3")
    flags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
    _progress(0.0, f"動画を読み込み中: {os.path.basename(video_path)}")
    try:
        completed = subprocess.run(
            [FFMPEG_PATH, "-y", "-i", video_path, "-vn", "-map", "0:a:0",
             "-c:a", "libmp3lame", "-b:a", "192k", audio_path],
            stdout=subprocess.DEVNULL, stderr=subprocess.PIPE,
            creationflags=flags)
        if completed.returncode != 0 or not os.path.exists(audio_path) \
                or os.path.getsize(audio_path) == 0:
            detail = completed.stderr.decode("utf-8", errors="replace").strip()
            detail = detail.splitlines()[-1] if detail else "音声トラックがありません"
            raise RuntimeError(f"動画から音声を抽出できませんでした: {detail}")
        if cancel is not None and cancel():
            raise PostProcessCancelled()

        snapshot_count = 0
        if capture_scenes:
            _progress(0.02, "動画の画面変化を検出中...")
            snapshot_count = _extract_video_snapshots(
                video_path, folder, interval=scene_interval,
                threshold=scene_threshold, max_edge=max_edge,
                jpeg_quality=jpeg_quality, progress=progress, cancel=cancel)
            _progress(0.03, f"動画から画像を保存しました: {snapshot_count}枚")

        with open(os.path.join(folder, "metadata.txt"), "w", encoding="utf-8") as f:
            f.write(f"MEETING_NAME={stem}\n")
            f.write(f"START_TIME_EPOCH={now.timestamp()}\n")
            f.write(f"START_TIME_STR={now.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("AUDIO_SOURCE=動画ファイル\n")
            f.write("MODE=video_import\n")
            f.write(f"SNAPSHOT_COUNT={snapshot_count}\n")
            f.write("MARKER_COUNT=0\n")
            f.write("AUDIO_FILE=audio_main.mp3\n")
            f.write(f"LANGUAGE={language}\n")
            f.write(f"SOURCE_VIDEO={video_path}\n")
    except Exception:
        shutil.rmtree(folder, ignore_errors=True)
        raise

    _progress(0.03, f"動画の音声抽出完了: {os.path.basename(audio_path)}")
    return folder


def post_process_folder(folder, progress=None, cancel=None):
    """Transcribe a meeting folder and generate the report files.

    progress: optional callable(fraction | None, message) for UI feedback.
              Post-processing a one-hour meeting takes tens of minutes on CPU,
              so every phase reports where it is.
    cancel:   optional callable() -> bool, polled at each checkpoint. When it
              returns True, PostProcessCancelled is raised.
    """
    report = _Reporter(progress, cancel)
    report(0.0, f"後処理開始: {os.path.basename(folder)}")
    data = collect_meeting_data(folder)
    report(0.02, f"音声: {data['audio_name']} / 画像: {len(data['images'])}枚")

    transcribe_meeting(data, report, span=(0.05, 0.80))

    report(0.82, "スライドOCR")
    try:
        ocr_images(data, report)
    except PostProcessCancelled:
        raise
    except Exception as e:
        report(None, f"[OCRエラー] {e}")

    report(0.88, "要約フェーズ")
    summarize_meeting(data, report)
    report(0.95, "議事録を出力中")

    md_path = os.path.join(folder, "meeting_report.md")
    render_markdown(data, md_path)

    cfg = load_app_settings()
    if cfg.get("export_html", True):
        try:
            html_path = os.path.join(folder, "meeting_report.html")
            _, size_mb = render_html(
                data, html_path,
                embed_audio=bool(cfg.get("html_embed_audio", False)),
                embed_images=bool(cfg.get("html_embed_images", True)))
            report(0.98, f"HTML を出力しました ({size_mb:.1f}MB)")
        except Exception as e:
            report(None, f"[HTML出力エラー] {e}")
            traceback.print_exc()
    if cfg.get("export_docx", False):
        try:
            render_docx(data, os.path.join(folder, "meeting_report.docx"))
            report(0.99, "DOCX を出力しました")
        except Exception as e:
            report(None, f"[DOCX出力エラー] {e}")

    report(1.0, f"議事録を出力しました: {os.path.basename(md_path)}")
    print(f"  - 発言 {len(data['lines'])}行 / 画像 {len(data['images'])}枚")
    return md_path


# ======================================================================

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="GijirokuStudio v2")
    parser.add_argument("--post-process", metavar="FOLDER",
        help="指定フォルダの音声を高精度文字起こしし、スクリーンショットと統合したMarkdownを出力")
    parser.add_argument("--import-video", metavar="VIDEO",
        help="動画から音声を抽出し、文字起こしと議事録を生成")
    parser.add_argument("--video-snapshots", action="store_true",
        help="動画入力時に画面変化を画像として保存し、議事録に表示")
    args = parser.parse_args()

    if args.import_video:
        try:
            imported_folder = import_video_file(
                args.import_video, capture_scenes=args.video_snapshots)
            post_process_folder(imported_folder)
            print(f"  - 保存先: {imported_folder}")
        except Exception as e:
            print(f"[ERROR] {e}")
            sys.exit(1)
    elif args.post_process:
        try:
            post_process_folder(args.post_process)
        except Exception as e:
            print(f"[ERROR] {e}")
            sys.exit(1)
    else:
        root = tk.Tk()
        app = MeetingRecorderGUI(root)

        def _on_close():
            if app.is_recording:
                app.is_recording = False
                app.stop_event.set()
                root.after(2000, root.destroy)
            else:
                root.destroy()

        root.protocol("WM_DELETE_WINDOW", _on_close)
        root.mainloop()
